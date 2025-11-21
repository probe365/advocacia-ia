from flask import Blueprint, jsonify, request, render_template, flash, redirect, url_for, send_from_directory, send_file, session, g
from flask_login import login_required
from app.services.cadastro_service import CadastroService
from cadastro_manager import CadastroManager
from pipeline import Pipeline
from werkzeug.utils import secure_filename
import os, hashlib, logging, base64
import re

logger = logging.getLogger(__name__)

def _compute_case_code(raw_id: str) -> str:
    """Retorna código curto estável (caso_<8hex>). Se já vier no formato, mantém para evitar re-hash."""
    try:
        if raw_id.startswith('caso_') and len(raw_id) >= 13:
            return raw_id  # já é um código
        return "caso_" + hashlib.sha1(raw_id.encode('utf-8')).hexdigest()[:8]
    except Exception:
        return f"caso_{raw_id[:8]}"

def _short_case_code(id_processo: str) -> str:
    try:
        return 'caso_' + hashlib.sha1(str(id_processo).encode('utf-8')).hexdigest()[:8]
    except Exception:
        return f'caso_{str(id_processo)[:8]}'

def _parse_firac_raw_text(raw_text: str) -> dict:
    """
    Parser robusto para converter FIRAC em formato markdown/texto para dicionário.
    
    Suporta formatos:
    - "**Fatos:**" ou "**Fatos**"
    - "1. **Fatos:**" (numerado)
    - Variações com acentos (Questão/Questao, Aplicação/Aplicacao, Conclusão/Conclusao)
    
    Returns:
        Dict com keys: facts, issue, rules, application, conclusion
    """
    if not raw_text or not raw_text.strip():
        logger.warning("[FIRAC PARSER] Raw text está vazio")
        return {'facts': '', 'issue': '', 'rules': '', 'application': '', 'conclusion': ''}
    
    # IMPROVED PARSER: Suporta formatos numerados e não numerados
    facts_match = re.search(
        r'(?:\d+\.\s+)?\*\*Fatos:?\*\*\s*(.*?)(?=\n\s*(?:\d+\.\s+)?\*\*(?:Quest[aã]o|Regras|Aplica[çc][aã]o|Conclus[aã]o)|\Z)', 
        raw_text, 
        re.DOTALL | re.IGNORECASE
    )
    issue_match = re.search(
        r'(?:\d+\.\s+)?\*\*Quest[aã]o:?\*\*\s*(.*?)(?=\n\s*(?:\d+\.\s+)?\*\*(?:Fatos|Regras|Aplica[çc][aã]o|Conclus[aã]o)|\Z)', 
        raw_text, 
        re.DOTALL | re.IGNORECASE
    )
    rules_match = re.search(
        r'(?:\d+\.\s+)?\*\*Regras:?\*\*\s*(.*?)(?=\n\s*(?:\d+\.\s+)?\*\*(?:Fatos|Quest[aã]o|Aplica[çc][aã]o|Conclus[aã]o)|\Z)', 
        raw_text, 
        re.DOTALL | re.IGNORECASE
    )
    app_match = re.search(
        r'(?:\d+\.\s+)?\*\*Aplica[çc][aã]o:?\*\*\s*(.*?)(?=\n\s*(?:\d+\.\s+)?\*\*(?:Fatos|Quest[aã]o|Regras|Conclus[aã]o)|\Z)', 
        raw_text, 
        re.DOTALL | re.IGNORECASE
    )
    concl_match = re.search(
        r'(?:\d+\.\s+)?\*\*Conclus[aã]o:?\*\*\s*(.*?)(?=\n\s*(?:\d+\.\s+)?\*\*|\Z)', 
        raw_text, 
        re.DOTALL | re.IGNORECASE
    )
    
    result = {
        'facts': facts_match.group(1).strip() if facts_match else '',
        'issue': issue_match.group(1).strip() if issue_match else '',
        'rules': rules_match.group(1).strip() if rules_match else '',
        'application': app_match.group(1).strip() if app_match else '',
        'conclusion': concl_match.group(1).strip() if concl_match else ''
    }
    
    # Log parsing results
    parsed_count = sum(1 for v in result.values() if v)
    logger.info(f"[FIRAC PARSER] Successfully parsed {parsed_count}/5 fields from raw text")
    
    if parsed_count == 0:
        logger.warning("[FIRAC PARSER] Nenhum campo foi parseado! Verifique formato do raw text")
    
    return result

processos_bp = Blueprint('processos', __name__, url_prefix='/processos')
service = CadastroService()

@processos_bp.route('/<id_cliente>/novo', methods=['GET','POST'])
@login_required
def novo_processo(id_cliente):
    if request.method == 'POST':
        dados = request.form.to_dict()
        try:
            proc_id = service.create_processo(id_cliente, dados)
            flash('Processo criado', 'success')
            return redirect(url_for('clientes.list_clientes_ui'))
        except Exception as e:
            flash(f'Erro: {e}', 'danger')
    advogados = service.list_advogados()
    return render_template('novo_processo.html', id_cliente=id_cliente, advogados=advogados)

# --- Bulk CSV Upload Routes ---
@processos_bp.route('/<id_cliente>/bulk-upload', methods=['GET'])
@login_required
def bulk_upload_form(id_cliente):
    """Exibe formulário para upload de múltiplos processos via CSV"""
    cliente = service.get_cliente(id_cliente)
    if not cliente:
        flash('Cliente não encontrado', 'danger')
        return redirect(url_for('clientes.list_clientes_ui'))
    return render_template('bulk_upload_processos.html', id_cliente=id_cliente, cliente=cliente)

@processos_bp.route('/api/<id_cliente>/bulk-upload', methods=['POST'])
@login_required
def bulk_upload_api(id_cliente):
    """API para upload de múltiplos processos via CSV (drag-and-drop)"""
    try:
        # Valida se cliente existe
        cliente = service.get_cliente(id_cliente)
        if not cliente:
            return jsonify({
                "status": "erro",
                "mensagem": "Cliente não encontrado"
            }), 404
        
        # Valida arquivo
        if 'csv_file' not in request.files:
            return jsonify({
                "status": "erro",
                "mensagem": "Nenhum arquivo enviado"
            }), 400
        
        csv_file = request.files['csv_file']
        if not csv_file or csv_file.filename == '':
            return jsonify({
                "status": "erro",
                "mensagem": "Arquivo vazio"
            }), 400
        
        # Valida extensão
        if not csv_file.filename.lower().endswith('.csv'):
            return jsonify({
                "status": "erro",
                "mensagem": "Arquivo deve ser CSV"
            }), 400
        
        # Lê conteúdo CSV
        try:
            csv_content = csv_file.read().decode('utf-8')
        except UnicodeDecodeError:
            return jsonify({
                "status": "erro",
                "mensagem": "Arquivo CSV não está em UTF-8"
            }), 400
        
        # Processa CSV via CadastroManager
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        resultado = mgr.bulk_create_processos_from_csv(id_cliente, csv_content)
        
        # Log para auditoria
        logger.info(f"Bulk upload CSV concluído para cliente {id_cliente}: {resultado['processos_criados']} processos criados")
        
        return jsonify(resultado), 200 if resultado['status'] == 'sucesso' else 400
        
    except Exception as e:
        logger.error(f"Erro no bulk upload API: {e}", exc_info=True)
        return jsonify({
            "status": "erro",
            "mensagem": str(e)
        }), 500

@processos_bp.route('/api/<id_cliente>/bulk-upload/preview', methods=['POST'])
@login_required
def bulk_upload_preview(id_cliente):
    """Preview dos processos antes de confirmar upload"""
    try:
        if 'csv_file' not in request.files:
            return jsonify({"status": "erro", "mensagem": "Nenhum arquivo enviado"}), 400
        
        csv_file = request.files['csv_file']
        
        try:
            csv_content = csv_file.read().decode('utf-8')
        except UnicodeDecodeError:
            return jsonify({"status": "erro", "mensagem": "Arquivo não está em UTF-8"}), 400
        
        # Parse CSV para preview
        import csv
        from io import StringIO
        
        reader = csv.DictReader(StringIO(csv_content))
        rows = list(reader)
        
        if not rows:
            return jsonify({"status": "erro", "mensagem": "CSV vazio"}), 400
        
        # Retorna preview (máximo 100 linhas)
        preview_rows = rows[:100]
        
        return jsonify({
            "status": "sucesso",
            "total_linhas": len(rows),
            "preview": preview_rows,
            "colunas": reader.fieldnames
        }), 200
        
    except Exception as e:
        logger.error(f"Erro no preview: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

@processos_bp.route('/api/<id_cliente>', methods=['GET'])
def list_processos_api(id_cliente):
    return jsonify(service.list_processos_do_cliente(id_cliente))

@processos_bp.route('/api/<id_cliente>', methods=['POST'])
def create_processo_api(id_cliente):
    dados = request.get_json() or {}
    if not dados.get('nome_caso'):
        return jsonify({'erro':'nome_caso requerido'}),400
    proc_id = service.create_processo(id_cliente, dados)
    dados['id_processo'] = proc_id
    return jsonify({'mensagem':'Processo criado','processo':dados}),201

@processos_bp.route('/api/delete/<id_processo>', methods=['DELETE'])
@login_required
def delete_processo_api(id_processo):
    """Deleta um processo (apenas para desenvolvimento/testes)"""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        
        # Verificar se processo existe
        processo = service.get_processo(id_processo)
        if not processo:
            return jsonify({
                "status": "erro",
                "mensagem": "Processo não encontrado"
            }), 404
        
        # Deletar processo
        success = mgr.delete_processo(id_processo)
        
        if success:
            logger.warning(f"⚠️ Processo DELETADO: {id_processo} (desenvolvimento)")
            return jsonify({
                "status": "sucesso",
                "mensagem": "Processo deletado com sucesso"
            }), 200
        else:
            return jsonify({
                "status": "erro",
                "mensagem": "Falha ao deletar processo"
            }), 500
            
    except Exception as e:
        logger.error(f"Erro ao deletar processo {id_processo}: {e}", exc_info=True)
        return jsonify({
            "status": "erro",
            "mensagem": str(e)
        }), 500

@processos_bp.route('/painel/<id_processo>')
@login_required
def painel_processo(id_processo):
    proc = service.get_processo(id_processo)
    if not proc:
        flash('Processo não encontrado','warning')
        return redirect(url_for('clientes.list_clientes_ui'))
    # Enriquecer com dados do advogado responsável (se houver)
    advogado = None
    try:
        advogado_oab = proc.get('advogado_oab') if isinstance(proc, dict) else getattr(proc, 'advogado_oab', None)
        if advogado_oab:
            advogado = service.get_advogado(advogado_oab)
    except Exception:
        advogado = None
    # Docs & pipeline context could be loaded here
    pipeline = Pipeline(case_id=id_processo)
    documentos = pipeline.list_unique_case_documents()
    chat_history = session.get(f'chat_history_{id_processo}', [])
    # Código curto estável para identificação visual (caso_<8 hex>)
    raw = str(proc.get('id_processo') if isinstance(proc, dict) else getattr(proc, 'id_processo', id_processo))
    # Usar o id original (já vem no formato caso_XXXXXXXX) para evitar divergência com URL.
    case_code = _compute_case_code(raw)
    advs = []
    try:
        advs = service.list_advogados()
    except Exception:
        advs = []
    return render_template('processo.html', processo=proc, documentos=documentos, chat_history=chat_history, case_code=case_code, advogado=advogado, advogados=advs)

@processos_bp.route('/ui/<id_processo>/atualizar_advogado', methods=['POST'])
@login_required
def ui_atualizar_advogado(id_processo):
    """Atualiza o advogado responsável via HTMX."""
    novo_oab = (request.form.get('advogado_oab') or '').strip()
    try:
        service.update_processo(id_processo, {'advogado_oab': novo_oab})
        adv = service.get_advogado(novo_oab) if novo_oab else None
        if adv:
            return f"<div class='alert alert-info py-2 small'><strong>Advogado Responsável:</strong> {adv['nome']} (OAB {adv['oab']}){' - Área: '+adv['area_atuacao'] if adv.get('area_atuacao') else ''}</div>"
        else:
            return "<div class='alert alert-info py-2 small'><em>Nenhum advogado associado a este processo.</em></div>"
    except Exception as e:
        return f"<div class='alert alert-danger py-2 small'>Erro ao atualizar: {e}</div>", 400

@processos_bp.route('/ui/<id_processo>/atualizar_tipo_parte', methods=['POST'])
@login_required
def ui_atualizar_tipo_parte(id_processo):
    """Atualiza o papel do cliente (tipo_parte) via HTMX."""
    from app.utils.tipo_parte_helpers import validate_tipo_parte, get_tipo_parte_label
    
    novo_tipo_parte = (request.form.get('tipo_parte') or '').strip().lower()
    
    # Validate tipo_parte
    if novo_tipo_parte and not validate_tipo_parte(novo_tipo_parte):
        return f"<div class='alert alert-danger py-2 small'>Erro: papel inválido. Valores válidos: autor, reu, terceiro, reclamante, reclamada</div>", 400
    
    try:
        # Set to None if empty string
        tipo_parte_value = novo_tipo_parte if novo_tipo_parte else None
        service.update_processo(id_processo, {'tipo_parte': tipo_parte_value})
        
        # Fetch updated processo to display
        processo = service.get_processo(id_processo)
        adv = service.get_advogado(processo.get('advogado_oab')) if processo.get('advogado_oab') else None
        
        # Build response HTML with both advogado and tipo_parte
        html = "<div class='alert alert-info py-2 small'>"
        if adv:
            html += f"<strong>Advogado Responsável:</strong> {adv['nome']} (OAB {adv['oab']}){' - Área: '+adv['area_atuacao'] if adv.get('area_atuacao') else ''}"
        else:
            html += "<em>Nenhum advogado associado a este processo.</em>"
        
        if tipo_parte_value:
            html += f"<br><strong>Papel do Cliente:</strong> <span class='badge bg-primary'>{tipo_parte_value.upper()}</span>"
        
        html += "</div>"
        return html
    except Exception as e:
        return f"<div class='alert alert-danger py-2 small'>Erro ao atualizar: {e}</div>", 400

# --- UI compatibility endpoints (previous /ui paths) ---
@processos_bp.route('/ui/<id_cliente>/processos', methods=['GET'])
def ui_list_processos_cliente(id_cliente):
    processos = service.list_processos_do_cliente(id_cliente)
    q = (request.args.get('q') or '').strip().lower()
    enriched = []
    for p in processos:
        # Mantém código original sem re-hash para consistência visual.
        code = _compute_case_code(p.get('id_processo'))
        p = dict(p)
        p['case_code'] = code
        enriched.append(p)
    if q:
        enriched = [p for p in enriched if q in (p.get('nome_caso') or '').lower() or q in (p.get('numero_cnj') or '').lower() or q in p['case_code'].lower()]
    processos = enriched
    return render_template('_lista_processos.html', processos=processos, id_cliente=id_cliente)

@processos_bp.route('/ui/<id_processo>/uploadform', methods=['GET'])
def ui_upload_form(id_processo):
    return render_template('_form_upload.html', id_processo=id_processo)

@processos_bp.route('/ui/<id_processo>/resumo', methods=['POST'])
def ui_resumo(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or request.values.get('focus') or '').strip()
        resumo, from_cache = pipeline.summarize_with_cache(focus or 'Resumo geral do caso')
        if not resumo or resumo.startswith("Sem conteúdo"):
            try:
                docs = pipeline.case_store.get(include=["documents"], limit=3)
                raw_texts = []
                if docs and docs.get('documents'):
                    for chunk_list in docs['documents'][:3]:
                        if isinstance(chunk_list, list):
                            raw_texts.extend(chunk_list)
                        else:
                            raw_texts.append(chunk_list)
                fallback_text = "\n\n".join(raw_texts)[:1200] or "Nenhum conteúdo disponível para resumo."
                resumo = f"(Fallback) Pré-visualização dos primeiros trechos:\n\n{fallback_text}"
            except Exception as fe:
                resumo = f"Não foi possível gerar fallback: {fe}"
        badge = "<span class='badge bg-info ms-2'>cache</span>" if from_cache else ""
        return f"""
        <div class='card mt-4'>
            <div class='card-header'>Resumo do Caso {badge}</div>
            <div class='card-body'><p style='white-space: pre-wrap;'>{resumo}</p>
              <div class='d-flex gap-2 flex-wrap'>
                <button class='btn btn-sm btn-outline-secondary'
                  hx-post='/processos/ui/{id_processo}/export/resumo'
                  hx-vals='js:{{focus: document.getElementById("focus-resumo").value}}'
                  hx-target='#download-area'
                  hx-swap='outerHTML'>Exportar TXT</button>
                <button class='btn btn-sm btn-outline-secondary'
                  hx-post='/processos/ui/{id_processo}/export/resumo/pdf'
                  hx-vals='js:{{focus: document.getElementById("focus-resumo").value}}'
                  hx-target='#download-area-pdf'
                  hx-swap='outerHTML'>Exportar PDF</button>
              </div>
              <div id='download-area' class='mt-2'></div>
              <div id='download-area-pdf' class='mt-1'></div>
            </div>
        </div>
        """
    except Exception as e:
        return f"<div class='alert alert-danger mt-4'>Erro: {e}</div>"

@processos_bp.route('/ui/<id_processo>/export/resumo', methods=['POST'])
def ui_export_resumo(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or '').strip()
        resumo = pipeline.summarize(query_for_relevance=focus or 'Resumo geral do caso')
        code = _compute_case_code(id_processo)
        filename = f"resumo_{code}.txt"
        import base64, html
        b64 = base64.b64encode(resumo.encode('utf-8')).decode('utf-8')
        safe_name = html.escape(filename)
        return f"<a download='{safe_name}' href='data:text/plain;base64,{b64}' class='btn btn-sm btn-success'>Baixar {safe_name}</a>"
    except Exception as e:
        return f"<div class='alert alert-danger mt-2'>Erro ao exportar: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/export/resumo/pdf', methods=['POST'])
def ui_export_resumo_pdf(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or '').strip()
        resumo = pipeline.summarize(query_for_relevance=focus or 'Resumo geral do caso')
        code = _compute_case_code(id_processo)
        filename = f"resumo_{code}.pdf"
        export_dir = pipeline.case_dir / 'exports'
        export_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = export_dir / filename
        generated = False
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_margins(15, 15, 15)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, f"Resumo do Caso - {code}", ln=1)
            pdf.set_font("Arial", size=11)
            # Tratar parágrafos preservando quebras em branco
            paragraphs = [p.strip() for p in resumo.split('\n')]
            for para in paragraphs:
                if not para:
                    pdf.ln(4)
                    continue
                # multi_cell faz wrap automático; substitui tabs múltiplos espaços
                pdf.multi_cell(0, 6, para.replace('\t', '    '))
                pdf.ln(1)
            pdf.output(str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 100:
                generated = True
            else:
                logger.warning('PDF resumo não gerado ou vazio', extra={'case_id': id_processo, 'code': code, 'path': str(pdf_path)})
        except ImportError:
            logger.info('fpdf não instalada - fallback para TXT')
        except Exception:
            logger.exception('Erro gerando PDF resumo')
        if generated and pdf_path.exists():
            try:
                size = pdf_path.stat().st_size
            except Exception:
                size = -1
            if size > 0:
                return f"<a class='btn btn-sm btn-success' href='/processos/ui/{id_processo}/download/exports/{filename}' target='_blank'>Baixar {filename} ({size//1024} KB)</a>"
            else:
                logger.error('Arquivo PDF resumo ausente ou vazio após geração', extra={'case_id': id_processo, 'code': code, 'path': str(pdf_path), 'size': size})
        b64 = base64.b64encode(resumo.encode('utf-8')).decode('utf-8')
        debug = ''
        if pdf_path and not pdf_path.exists():
            debug = f" (debug: arquivo não encontrado em {pdf_path})"
        return f"<a download='{filename.replace('.pdf','.txt')}' href='data:text/plain;base64,{b64}' class='btn btn-sm btn-warning'>Baixar como TXT (fallback)</a><small class='text-muted ms-2'>PDF não gerado{debug}</small>"
    except Exception as e:
        return f"<div class='alert alert-danger mt-2'>Erro ao exportar PDF: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/download/exports/<path:filename>')
def ui_download_export(id_processo, filename):
    """Serve arquivos de export (PDF/TXT) com diagnóstico detalhado em caso de falha."""
    try:
        pipeline = Pipeline(case_id=id_processo)
        export_dir = pipeline.case_dir / 'exports'
        if not export_dir.exists():
            logger.error('export_dir inexistente', extra={'case_id': id_processo, 'path': str(export_dir)})
            return f"<div class='alert alert-warning mt-2'>Pasta de export não existe.</div>", 404
        target = (export_dir / filename).resolve()
        # Segurança básica: garantir que continua dentro de export_dir
        if not str(target).startswith(str(export_dir.resolve())):
            return "Path inválido", 400
        if not target.exists():
            logger.warning('arquivo export não encontrado', extra={'case_id': id_processo, 'file': str(target)})
            return f"<div class='alert alert-warning mt-2'>Arquivo não encontrado: {filename}</div>", 404
        mime = 'application/pdf' if filename.lower().endswith('.pdf') else 'text/plain'
        try:
            return send_file(str(target), mimetype=mime, as_attachment=True, download_name=filename)
        except TypeError:
            # Para compatibilidade Flask mais antiga sem download_name
            return send_file(str(target), mimetype=mime, as_attachment=True)
    except Exception as e:
        logger.exception('Erro download export')
        return f"<div class='alert alert-danger mt-2'>Erro download: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/analise/firac', methods=['POST'])
def ui_analise_firac(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or '').strip()
        result = pipeline.generate_firac(focus=focus)
        if result.get('data'):
            d = result['data']
            # Render FIRAC formatado em HTML
            cached_badge = "<span class='badge bg-info ms-2'>cache</span>" if result.get('cached') else ''
            html = [f"<div class='firac-result'><div class='mb-2 small text-muted'>FIRAC {'(cache reutilizado)' if result.get('cached') else '(gerado)'}{cached_badge}</div>"]
            facts = d.get('facts') or []
            html.append("<h6>Fatos (Facts)</h6><ol>" + "".join([f"<li>{f}</li>" for f in facts]) + "</ol>")
            html.append(f"<h6>Questão (Issue)</h6><p>{d.get('issue','')}</p>")
            rules = d.get('rules') or []
            html.append("<h6>Regras (Rule)</h6><ol>" + "".join([f"<li>{r}</li>" for r in rules]) + "</ol>")
            html.append(f"<h6>Aplicação (Application)</h6><p>{d.get('application','')}</p>")
            html.append(f"<h6>Conclusão (Conclusion)</h6><p>{d.get('conclusion','')}</p>")
            html.append("</div>")
            return "".join(html)
        else:
            return f"<pre style='white-space:pre-wrap;font-size:0.85rem;'>{result.get('raw')}</pre>"
    except Exception as e:
        return f"<div class='alert alert-danger'>Erro FIRAC: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/analise/firac/export/pdf', methods=['POST'])
def ui_export_firac_pdf(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or '').strip()
        result = pipeline.generate_firac(focus=focus)
        data = result.get('data')
        raw = result.get('raw')
        from datetime import datetime
        code = _compute_case_code(id_processo)
        filename = f"firac_{code}.pdf"
        export_dir = pipeline.case_dir / 'exports'
        export_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = export_dir / filename
        generated = False
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=12)
            pdf.add_page()
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, f"Análise FIRAC - {code}", ln=1)
            pdf.set_font("Arial", size=9)
            pdf.cell(0, 6, f"Gerado: {datetime.utcnow().isoformat()}Z", ln=1)
            pdf.ln(2)
            def write_section(title, content_lines):
                pdf.set_font("Arial", "B", 11)
                pdf.cell(0, 7, title, ln=1)
                pdf.set_font("Arial", size=10)
                if isinstance(content_lines, list):
                    for i, l in enumerate(content_lines, 1):
                        pdf.multi_cell(0, 5, f"{i}. {l}")
                else:
                    for para in str(content_lines).split('\n'):
                        pdf.multi_cell(0, 5, para)
                pdf.ln(2)
            if data:
                write_section('Fatos', data.get('facts') or [])
                write_section('Questão', data.get('issue',''))
                write_section('Regras', data.get('rules') or [])
                write_section('Aplicação', data.get('application',''))
                write_section('Conclusão', data.get('conclusion',''))
            else:
                write_section('FIRAC (Texto)', raw)
            pdf.output(str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 100:
                generated = True
            else:
                logger.warning('PDF FIRAC não gerado ou vazio', extra={'case_id': id_processo, 'code': code, 'path': str(pdf_path)})
        except ImportError:
            logger.info('fpdf não instalada - fallback para TXT FIRAC')
        except Exception:
            logger.exception('Erro gerando PDF FIRAC')
        if generated and pdf_path.exists():
            try:
                size = pdf_path.stat().st_size
            except Exception:
                size = -1
            if size > 0:
                return f"<a class='btn btn-sm btn-success' href='/processos/ui/{id_processo}/download/exports/{filename}' target='_blank'>Baixar {filename} ({size//1024} KB)</a>"
            else:
                logger.error('Arquivo PDF FIRAC ausente ou vazio após geração', extra={'case_id': id_processo, 'code': code, 'path': str(pdf_path), 'size': size})
        b64 = base64.b64encode((raw or 'FIRAC não disponível').encode('utf-8')).decode('utf-8')
        debug = ''
        if pdf_path and not pdf_path.exists():
            debug = f" (debug: arquivo não encontrado em {pdf_path})"
        return f"<a download='{filename.replace('.pdf','.txt')}' href='data:text/plain;base64,{b64}' class='btn btn-sm btn-warning'>Baixar TXT (fallback)</a><small class='text-muted ms-2'>PDF não gerado{debug}</small>"
    except Exception as e:
        return f"<div class='alert alert-danger'>Erro exportar FIRAC: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/peticao/form', methods=['GET'])
def ui_peticao_form(id_processo):
        """Retorna formulário simples (HTMX) para gerar petição com dados básicos de juízo/partes."""
        try:
                proc = service.get_processo(id_processo) or {}
                cliente = None
                advogado = None
                try:
                        if proc.get('id_cliente'):
                                cliente = service.get_cliente(proc.get('id_cliente'))
                except Exception:
                        cliente = None
                try:
                        if proc.get('advogado_oab'):
                                advogado = service.get_advogado(proc.get('advogado_oab'))
                except Exception:
                        advogado = None
                autor_nome = (cliente or {}).get('nome_completo','') or (cliente or {}).get('nome_completo', '') or ''
                advogado_label = ''
                if advogado:
                        advogado_label = f"<div class='alert alert-secondary p-1 mb-2 small'>Advogado: {advogado.get('nome')} (OAB {advogado.get('oab')})</div>"
                return f"""
                {advogado_label}
                <form hx-post='/processos/ui/{id_processo}/peticao/gerar' hx-target='#peticao-result' hx-swap='innerHTML'>
                    <div class='row g-2'>
                        <div class='col-md-6'>
                            <label class='form-label small mb-0'>Comarca</label>
                            <input name='juizo_comarca' class='form-control form-control-sm' placeholder='Ex: São Paulo'>
                        </div>
                        <div class='col-md-2'>
                            <label class='form-label small mb-0'>UF</label>
                            <input name='juizo_uf' class='form-control form-control-sm' maxlength='2' placeholder='SP'>
                        </div>
                        <div class='col-md-2'>
                            <label class='form-label small mb-0'>Vara</label>
                            <input name='juizo_vara' class='form-control form-control-sm' placeholder='1ª'>
                        </div>
                        <div class='col-md-2'>
                            <label class='form-label small mb-0'>Esp.</label>
                            <input name='juizo_especialidade' class='form-control form-control-sm' placeholder='Cível' value='Cível'>
                        </div>
                        <div class='col-md-6'>
                            <label class='form-label small mb-0'>Autor</label>
                            <input name='autor_nome' class='form-control form-control-sm' value='{autor_nome}' readonly>
                        </div>
                        <div class='col-md-6'>
                            <label class='form-label small mb-0'>Réu (nome completo)</label>
                            <input name='reu_nome' class='form-control form-control-sm'>
                        </div>
                        <div class='col-md-4'>
                            <label class='form-label small mb-0'>Valor da Causa</label>
                            <input name='valor_causa_num' class='form-control form-control-sm' placeholder='10000,00'>
                        </div>
                        <div class='col-md-4'>
                            <label class='form-label small mb-0'>Valor por Extenso</label>
                            <input name='valor_causa_ext' class='form-control form-control-sm' placeholder='dez mil reais'>
                        </div>
                        <div class='col-md-4'>
                            <label class='form-label small mb-0'>Provas / Observ.</label>
                            <input name='texto_provas_especificas' class='form-control form-control-sm' placeholder='(opcional)'>
                        </div>
                        <div class='col-12 d-flex justify-content-end'>
                            <button class='btn btn-sm btn-primary'>Gerar Petição</button>
                        </div>
                    </div>
                </form>
                """
        except Exception as e:
                return f"<div class='alert alert-danger'>Erro ao montar formulário: {e}</div>"

@processos_bp.route('/ui/<id_processo>/peticao/gerar', methods=['POST'])
def ui_peticao_gerar(id_processo):
    """Gera rascunho de petição inicial usando FIRAC + inputs. Retorna bloco HTML com texto e export options."""
    try:
        pipeline = Pipeline(case_id=id_processo)
        proc = service.get_processo(id_processo) or {}
        cliente = service.get_cliente(proc.get('id_cliente')) if proc.get('id_cliente') else None
        advogado = service.get_advogado(proc.get('advogado_oab')) if proc.get('advogado_oab') else None
        firac = pipeline.generate_firac()
        
        # DEBUG: Log what we got from generate_firac
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"[PETITION DEBUG] FIRAC returned: {firac}")
        
        data_firac = firac.get('data') or {}
        logger.info(f"[PETITION DEBUG] data_firac extracted: {data_firac}")
        
        # If data is None/empty but we have raw text, use the improved parser
        if not data_firac and firac.get('raw'):
            raw = firac.get('raw', '')
            logger.info(f"[PETITION DEBUG] Parsing raw FIRAC text using improved parser")
            data_firac = _parse_firac_raw_text(raw)
            logger.info(f"[PETITION DEBUG] Parsed data_firac: {data_firac}")
        
        # Robustez: tentar extrair conclusão se vier vazia ou ausente
        if (not data_firac) or (not data_firac.get('facts') or not data_firac.get('conclusion')):
            logger.warning(f"[PETITION DEBUG] FIRAC data missing or incomplete, trying to parse raw")
            raw_firac = firac.get('raw', '')
            import re
            concl = ''
            # Tenta padrão JSON-like ou texto marcado
            m = re.search(r'"?conclusion"?\s*[:=]\s*"([^"\n]+)"', raw_firac, re.IGNORECASE)
            if m:
                concl = m.group(1).strip()
            else:
                m2 = re.search(r'Conclus[aã]o\s*[:\-]\s*(.+?)(?:\n\s*\n|$)', raw_firac, re.IGNORECASE | re.DOTALL)
                if m2:
                    concl = m2.group(1).strip()
            if concl and not data_firac.get('conclusion'):
                data_firac['conclusion'] = concl
        
        logger.info(f"[PETITION DEBUG] Final data_firac to use: {data_firac}")
        form = request.form
        adv_name = (advogado or {}).get('nome') or ''
        adv_oab_raw = (advogado or {}).get('oab') or ''
        oab_uf = 'XX'; oab_num = adv_oab_raw
        import re
        m = re.match(r'([A-Za-z]{2})\s*-?\s*(.*)', adv_oab_raw)
        if m:
            oab_uf = m.group(1).upper(); oab_num = m.group(2).strip()
        dados_ui = {
            'juizo': {
                'vara': form.get('juizo_vara') or '',
                'especialidade': form.get('juizo_especialidade') or 'CÍVEL',
                'comarca': form.get('juizo_comarca') or '',
                'uf': (form.get('juizo_uf') or '').upper() or 'XX'
            },
            'autor': {'nome_completo_ou_razao_social': (cliente or {}).get('nome_completo') or form.get('autor_nome') or 'AUTOR'},
            'reu': {'nome': form.get('reu_nome') or 'RÉU'},
            'advogado': {
                'nome': adv_name,
                'oab_uf': oab_uf,
                'oab_numero': oab_num,
                'email': (advogado or {}).get('email')
            },
            'outros': {
                'valor_causa_num': form.get('valor_causa_num') or '',
                'valor_causa_ext': form.get('valor_causa_ext') or '',
                'texto_provas_especificas': form.get('texto_provas_especificas') or ''
            }
        }
        # Convert FIRAC lists to strings for petition generation
        facts_raw = data_firac.get('facts', [])
        rules_raw = data_firac.get('rules', [])
        
        facts_str = '\n'.join(facts_raw) if isinstance(facts_raw, list) else str(facts_raw)
        rules_str = '\n'.join(rules_raw) if isinstance(rules_raw, list) else str(rules_raw)
        
        peticao_txt = pipeline.generate_peticao_rascunho(dados_ui, {
            'facts': facts_str,
            'issue': data_firac.get('issue', ''),
            'rules': rules_str,
            'application': data_firac.get('application', ''),
            'conclusion': data_firac.get('conclusion', '')
        })
        # Sanitização de mensagens de ausência de dados para evitar texto de desculpas
        for marker in ["Desculpe", "Por favor, forneça", "não forneceu a conclusão"]:
            if marker in peticao_txt:
                peticao_txt = peticao_txt.replace(marker, '')
        safe = peticao_txt.replace('\n', '<br>').replace('  ', '&nbsp;&nbsp;')
        return f"<div class='mt-3'><div class='d-flex justify-content-between align-items-center flex-wrap gap-2'><h6 class='mb-2'>Rascunho da Petição</h6><div class='btn-group btn-group-sm'><button class='btn btn-outline-secondary' hx-get='/processos/ui/{id_processo}/peticao/export/pdf' hx-target='#peticao-export-area' hx-swap='innerHTML'>Exportar PDF</button><button class='btn btn-outline-secondary' hx-get='/processos/ui/{id_processo}/peticao/export/docx' hx-target='#peticao-export-area' hx-swap='innerHTML'>Exportar DOCX</button></div></div><div class='border rounded p-2 bg-body-tertiary' style='max-height:55vh;overflow:auto;font-size:0.8rem;font-family:monospace;'>{safe}</div><div id='peticao-export-area' class='mt-2 small text-muted'></div></div>"
    except Exception as e:
        return f"<div class='alert alert-danger mt-2'>Erro ao gerar petição: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/peticao/export/pdf', methods=['GET'])
def ui_peticao_export_pdf(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        proc = service.get_processo(id_processo) or {}
        cliente = service.get_cliente(proc.get('id_cliente')) if proc.get('id_cliente') else None
        advogado = service.get_advogado(proc.get('advogado_oab')) if proc.get('advogado_oab') else None
        firac = pipeline.generate_firac()
        
        data_firac = firac.get('data') or {}
        
        # Use improved parser if data is empty but raw exists
        if not data_firac and firac.get('raw'):
            logger.info("[PDF EXPORT] Parsing raw FIRAC text")
            data_firac = _parse_firac_raw_text(firac.get('raw', ''))
        
        # Legacy fallback for conclusion only (mantido para compatibilidade)
        if (not data_firac) or (not data_firac.get('conclusion')):
            raw_firac = firac.get('raw', '')
            m2 = re.search(r'Conclus[aã]o\s*[:\-]\s*(.+?)(?:\n\s*\n|$)', raw_firac, re.IGNORECASE | re.DOTALL)
            if m2:
                concl = m2.group(1).strip()
                if not data_firac:
                    data_firac = {'facts': '', 'issue': '', 'rules': '', 'application': '', 'conclusion': concl}
                else:
                    data_firac['conclusion'] = data_firac.get('conclusion') or concl
        adv_name = (advogado or {}).get('nome') or ''
        adv_oab_raw = (advogado or {}).get('oab') or ''
        import re
        oab_uf='XX'; oab_num=adv_oab_raw
        m=re.match(r'([A-Za-z]{2})\s*-?\s*(.*)', adv_oab_raw)
        if m: oab_uf=m.group(1).upper(); oab_num=m.group(2).strip()
        dados_ui = {'juizo': {}, 'autor': {'nome_completo_ou_razao_social': (cliente or {}).get('nome_completo','AUTOR')}, 'reu': {'nome': 'RÉU'}, 'advogado': {'nome': adv_name, 'oab_uf': oab_uf, 'oab_numero': oab_num}, 'outros': {}}
        
        # Convert FIRAC lists to strings for petition generation (PDF export)
        facts_raw = data_firac.get('facts', [])
        rules_raw = data_firac.get('rules', [])
        facts_str = '\n'.join(facts_raw) if isinstance(facts_raw, list) else str(facts_raw)
        rules_str = '\n'.join(rules_raw) if isinstance(rules_raw, list) else str(rules_raw)
        
        peticao_txt = pipeline.generate_peticao_rascunho(dados_ui, {
            'facts': facts_str,
            'issue': data_firac.get('issue', ''),
            'rules': rules_str,
            'application': data_firac.get('application', ''),
            'conclusion': data_firac.get('conclusion', '')
        })
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_margins(15, 15, 15)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font('Arial', 'B', 13)
        pdf.multi_cell(0, 8, f'Petição Inicial - {id_processo}')
        pdf.ln(2)
        pdf.set_font('Arial', size=10)
        for para in peticao_txt.split('\n'):
            if not para.strip():
                pdf.ln(4)
                continue
            pdf.multi_cell(0, 5, para)
            pdf.ln(1)
        export_dir = pipeline.case_dir / 'exports'
        export_dir.mkdir(parents=True, exist_ok=True)
        filename = f'peticao_{_compute_case_code(id_processo)}.pdf'
        pdf_path = export_dir / filename
        pdf.output(str(pdf_path))
        if pdf_path.exists():
            return f"<a class='btn btn-sm btn-success' href='/processos/ui/{id_processo}/download/exports/{filename}' target='_blank'>Baixar {filename}</a>"
        return "<div class='text-danger'>Falha ao gerar PDF.</div>"
    except Exception as e:
        return f"<div class='alert alert-danger'>Erro exportar PDF petição: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/peticao/export/docx', methods=['GET'])
def ui_peticao_export_docx(id_processo):
    try:
        try:
            from docx import Document
        except ImportError:
            return "<div class='text-danger'>Dependência 'python-docx' não instalada. Instale para exportar DOCX: pip install python-docx</div>"
        pipeline = Pipeline(case_id=id_processo)
        proc = service.get_processo(id_processo) or {}
        cliente = service.get_cliente(proc.get('id_cliente')) if proc.get('id_cliente') else None
        advogado = service.get_advogado(proc.get('advogado_oab')) if proc.get('advogado_oab') else None
        firac = pipeline.generate_firac()
        
        data_firac = firac.get('data') or {}
        
        # Use improved parser if data is empty but raw exists
        if not data_firac and firac.get('raw'):
            logger.info("[DOCX EXPORT] Parsing raw FIRAC text")
            data_firac = _parse_firac_raw_text(firac.get('raw', ''))
        
        # Legacy fallback for conclusion only (mantido para compatibilidade)
        if (not data_firac) or (not data_firac.get('conclusion')):
            raw_firac = firac.get('raw', '')
            m2 = re.search(r'Conclus[aã]o\s*[:\-]\s*(.+?)(?:\n\s*\n|$)', raw_firac, re.IGNORECASE | re.DOTALL)
            if m2:
                concl = m2.group(1).strip()
                if not data_firac:
                    data_firac = {'facts': '', 'issue': '', 'rules': '', 'application': '', 'conclusion': concl}
                else:
                    data_firac['conclusion'] = data_firac.get('conclusion') or concl
        adv_name = (advogado or {}).get('nome') or ''
        adv_oab_raw = (advogado or {}).get('oab') or ''
        oab_uf='XX'; oab_num=adv_oab_raw
        m=re.match(r'([A-Za-z]{2})\s*-?\s*(.*)', adv_oab_raw)
        if m: oab_uf=m.group(1).upper(); oab_num=m.group(2).strip()
        dados_ui = {'juizo': {}, 'autor': {'nome_completo_ou_razao_social': (cliente or {}).get('nome_completo','AUTOR')}, 'reu': {'nome': 'RÉU'}, 'advogado': {'nome': adv_name, 'oab_uf': oab_uf, 'oab_numero': oab_num}, 'outros': {}}
        
        # Convert FIRAC lists to strings for petition generation (DOCX export)
        facts_raw = data_firac.get('facts', [])
        rules_raw = data_firac.get('rules', [])
        facts_str = '\n'.join(facts_raw) if isinstance(facts_raw, list) else str(facts_raw)
        rules_str = '\n'.join(rules_raw) if isinstance(rules_raw, list) else str(rules_raw)
        
        peticao_txt = pipeline.generate_peticao_rascunho(dados_ui, {
            'facts': facts_str,
            'issue': data_firac.get('issue', ''),
            'rules': rules_str,
            'application': data_firac.get('application', ''),
            'conclusion': data_firac.get('conclusion', '')
        })
        doc = Document()
        doc.add_heading(f'Petição Inicial - {id_processo}', level=1)
        for para in peticao_txt.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
            else:
                doc.add_paragraph('')
        export_dir = pipeline.case_dir / 'exports'
        export_dir.mkdir(parents=True, exist_ok=True)
        filename = f'peticao_{_compute_case_code(id_processo)}.docx'
        doc_path = export_dir / filename
        doc.save(str(doc_path))
        if doc_path.exists():
            return f"<a class='btn btn-sm btn-success' href='/processos/ui/{id_processo}/download/exports/{filename}' target='_blank'>Baixar {filename}</a>"
        return "<div class='text-danger'>Falha ao gerar DOCX.</div>"
    except Exception as e:
        return f"<div class='alert alert-danger'>Erro exportar DOCX petição: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/analise/riscos', methods=['POST'])
def ui_analise_riscos(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or '').strip()
        txt = pipeline.identify_legal_risks(focus=focus)
        return f"<pre style='white-space:pre-wrap;font-size:0.85rem;'>{txt}</pre>"
    except Exception as e:
        return f"<div class='alert alert-danger'>Erro riscos: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/analise/proximos_passos', methods=['POST'])
def ui_analise_proximos_passos(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        focus = (request.form.get('focus') or '').strip()
        txt = pipeline.suggest_next_steps(focus=focus)
        return f"<pre style='white-space:pre-wrap;font-size:0.85rem;'>{txt}</pre>"
    except Exception as e:
        return f"<div class='alert alert-danger'>Erro próximos passos: {e}</div>", 500

# ----------------- Documentos (upload/delete) -----------------

@processos_bp.route('/ui/<id_processo>/documentos/novo', methods=['POST'])
def ui_upload_documento(id_processo):
    try:
        pipeline = Pipeline(case_id=id_processo)
        f = request.files.get('file')
        if not f:
            return "<div class='alert alert-danger mt-2'>Arquivo não enviado.</div>"
        filename = secure_filename(f.filename)
        content = f.read()
        ext = os.path.splitext(filename)[1].lower()
        added_type = None
        ocr_info = ''
        # Salva cópia física para possibilitar preview (ex: imagens)
        try:
            uploads_dir = pipeline.case_dir / 'uploads'
            uploads_dir.mkdir(parents=True, exist_ok=True)
            with open(uploads_dir / filename, 'wb') as out_f:
                out_f.write(content)
        except Exception as fs_e:
            flash(f"Falha ao salvar arquivo para preview: {fs_e}", 'warning')
        if ext == '.pdf':
            txt = pipeline.ingestion_handler.add_pdf(content, source_name=filename)
            if not txt: ocr_info = " (sem texto extraível)"
            added_type = 'pdf'
        elif ext in ['.jpg', '.jpeg', '.png']:
            txt = pipeline.ingestion_handler.add_image(content, source_name=filename)
            if not txt: ocr_info = " (OCR vazio; placeholder adicionado)"
            added_type = 'image'
        elif ext == '.txt':
            # Suporte a texto simples
            try:
                try:
                    text_decoded = content.decode('utf-8')
                except UnicodeDecodeError:
                    text_decoded = content.decode('latin-1')
                if not text_decoded.strip():
                    return "<div class='alert alert-warning mt-2'>Arquivo .txt vazio.</div>"
                pipeline.ingestion_handler.add_text_direct(text_decoded, source_name=filename, metadata_override={"type": "text"})
                added_type = 'text'
            except Exception as de:
                return f"<div class='alert alert-danger mt-2'>Falha ao ler TXT: {de}</div>"
        elif ext in ['.mp3', '.wav']:
            pipeline.ingestion_handler.add_audio(content, source_name=filename, audio_format_suffix=ext, openai_client=pipeline.openai_client)
            added_type = 'audio'
        elif ext in ['.mp4', '.mov']:
            pipeline.ingestion_handler.add_video(content, source_name=filename, video_format_suffix=ext, openai_client=pipeline.openai_client)
            added_type = 'video'
        else:
            return f"<div class='alert alert-warning mt-2'>Tipo de arquivo não suportado: {ext}</div>"
        documentos = pipeline.list_unique_case_documents()
        lista_html = render_template('_lista_documentos.html', documentos=documentos)
        msg = f"<div class='alert alert-success mb-2 p-2'>Arquivo '{filename}' ({added_type}) processado{ocr_info}.</div>"
        response_html = msg + lista_html
        # Fecha modal e atualiza contagem via evento custom
        response_headers = {'HX-Trigger': 'fecharUploadModal'}
        try:
            logger.info('upload_documento', extra={'case_id': id_processo, 'case_code': _short_case_code(id_processo), 'filename': filename, 'type': added_type})
        except Exception:
            pass
        return response_html, 200, response_headers
    except Exception as e:
        return f"<div class='alert alert-danger mt-2'>Erro no upload: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/documentos/<filename>', methods=['DELETE'])
def ui_delete_documento(id_processo, filename):
    try:
        pipeline = Pipeline(case_id=id_processo)
        ok = pipeline.delete_document_by_filename(filename)
        documentos = pipeline.list_unique_case_documents()
        html = render_template('_lista_documentos.html', documentos=documentos)
        return html if ok else ("<div class='alert alert-danger mt-2'>Falha ao remover.</div>", 500)
    except Exception as e:
        return f"<div class='alert alert-danger mt-2'>Erro ao deletar: {e}</div>", 500

# ----------------- Chat do Caso -----------------
@processos_bp.route('/ui/<id_processo>/chat', methods=['POST'])
def ui_chat(id_processo):
    """Processa uma pergunta do usuário sobre o caso via formulário HTMX.
    Evita 415 Unsupported Media Type aceitando form-urlencoded ou multipart.
    """
    try:
        pipeline = Pipeline(case_id=id_processo)
        query = (request.form.get('query') or '').strip()
        scope = (request.form.get('scope') or 'case').lower()
        if not query:
            return "<div class='alert alert-warning p-2 m-2'>Pergunta vazia.</div>"
        history = session.get(f'chat_history_{id_processo}', [])
        response = pipeline.chat(query, history, search_scope=scope)
        answer = response.get('output', response.get('output_text')) or response.get('final_answer')
        if not answer:
            answer = response.get('answer') or str(response)
        history.append({'role':'user','content':query})
        history.append({'role':'assistant','content':answer})
        session[f'chat_history_{id_processo}'] = history
        try:
            logger.info('chat_turn', extra={'case_id': id_processo, 'case_code': _short_case_code(id_processo), 'question': query, 'answer_len': len(answer)})
        except Exception:
            pass
        return render_template('_conversa_chat.html', chat_history=history)
    except Exception as e:
        return f"<div class='alert alert-danger p-2 m-2'>Erro no chat: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/chat/clear', methods=['POST'])
def ui_chat_clear(id_processo):
    """Limpa o histórico do chat para o processo."""
    try:
        # Clear the session history
        session.pop(f'chat_history_{id_processo}', None)
        logger.info('chat_cleared', extra={'case_id': id_processo, 'case_code': _short_case_code(id_processo)})
        # Return empty chat container
        return render_template('_conversa_chat.html', chat_history=[])
    except Exception as e:
        return f"<div class='alert alert-danger p-2 m-2'>Erro ao limpar chat: {e}</div>", 500

@processos_bp.route('/ui/<id_processo>/arquivo/<path:filename>', methods=['GET'])
def ui_serve_uploaded_file(id_processo, filename):
    """Serve arquivos salvos do caso (imagens para preview)."""
    # Evita custo de recriar pipeline completo apenas para servir arquivo
    from pathlib import Path
    uploads_dir = Path('./cases') / id_processo / 'uploads'
    target = uploads_dir / filename
    if not target.exists():
        return "Arquivo não encontrado", 404
    # Segurança básica: garantir que o arquivo está dentro do diretório
    try:
        resp = send_from_directory(uploads_dir, filename)
        resp.headers['Cache-Control'] = 'max-age=300'
        return resp
    except Exception as e:
        return f"Erro ao servir arquivo: {e}", 500

# ============================================================
# FEATURE 1: Download CSV Template (Item 4 - DIA 2)
# ============================================================
@processos_bp.route('/api/<id_cliente>/bulk-upload/template', methods=['GET'])
@login_required
def bulk_upload_template(id_cliente):
    """Download do template CSV com TODOS os 17 campos disponíveis (Item 1 + básicos)."""
    try:
        import csv
        from io import StringIO
        
        # Validar cliente
        cliente = service.get_cliente(id_cliente)
        if not cliente:
            return jsonify({"status": "erro", "mensagem": "Cliente não encontrado"}), 404
        
        # Define todas as colunas disponíveis (5 básicos + 12 do Item 1)
        colunas = [
            "nome_caso",           # Obrigatório
            "numero_cnj",
            "status",
            "advogado_oab",
            "tipo_parte",
            "comarca",
            "vara",
            "juiz_nome",
            "data_distribuicao",
            "data_citacao",
            "data_audiencia",
            "valor_causa",
            "valor_condenacao",
            "tipo_acao",
            "grau_jurisdicao",
            "instancia",
            "observacoes"
        ]
        
        # Criar CSV com exemplo completo
        output = StringIO()
        writer = csv.DictWriter(output, fieldnames=colunas)
        
        writer.writeheader()
        writer.writerow({
            "nome_caso": "Ação de Cobrança Exemplo",
            "numero_cnj": "1234567-89.2023.8.26.0100",
            "status": "ATIVO",
            "advogado_oab": "SP123456",
            "tipo_parte": "autor",
            "comarca": "São Paulo",
            "vara": "1ª Vara Cível",
            "juiz_nome": "Dr. José Silva",
            "data_distribuicao": "2023-05-15",
            "data_citacao": "2023-06-20",
            "data_audiencia": "2023-08-10",
            "valor_causa": "15000.50",
            "valor_condenacao": "12000.00",
            "tipo_acao": "Cobrança",
            "grau_jurisdicao": "1º Grau",
            "instancia": "Primeira Instância",
            "observacoes": "Processo aguardando sentença"
        })
        
        csv_content = output.getvalue()
        
        # Gerar nome do arquivo
        filename = f"template_processos_{cliente.get('nome_completo', 'cliente').replace(' ', '_')}.csv"
        
        logger.info(f"Template CSV (17 campos) baixado para cliente {id_cliente}")
        
        # Converter para BytesIO para compatibilidade com send_file
        from io import BytesIO
        bytes_output = BytesIO()
        bytes_output.write(csv_content.encode('utf-8'))
        bytes_output.seek(0)
        
        return send_file(
            bytes_output,
            as_attachment=True,
            download_name=filename,
            mimetype='text/csv; charset=utf-8'
        )
    except Exception as e:
        logger.error(f"Erro ao gerar template: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

# ============================================================
# FEATURE 2: Bulk Upload History/Log
# ============================================================
@processos_bp.route('/api/<id_cliente>/bulk-upload/history', methods=['GET'])
@login_required
def bulk_upload_history(id_cliente):
    """Retorna histórico de uploads em massa para um cliente."""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        
        # Validar cliente
        cliente = service.get_cliente(id_cliente)
        if not cliente:
            return jsonify({"status": "erro", "mensagem": "Cliente não encontrado"}), 404
        
        # Obter histórico (armazenado em JSON na coluna de logs do cliente)
        # Para simplificar, vamos criar uma view que consolida uploads
        import json
        from datetime import datetime
        
        # Se não temos tabela de auditoria, criar estrutura simples
        # Aqui você pode conectar a um log real se tiver tabela de auditoria
        
        return jsonify({
            "status": "sucesso",
            "cliente_id": id_cliente,
            "historico": []  # Placeholder - integrar com auditoria real
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao obter histórico: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

# ============================================================
# FEATURE 3: Email Notification After Bulk Upload
# ============================================================
@processos_bp.route('/api/<id_cliente>/bulk-upload/notify', methods=['POST'])
@login_required
def bulk_upload_notify(id_cliente):
    """Envia notificação por email após bulk upload."""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from datetime import datetime
        import json
        
        # Dados do request
        data = request.get_json() or {}
        
        # Validar cliente
        cliente = service.get_cliente(id_cliente)
        if not cliente:
            return jsonify({"status": "erro", "mensagem": "Cliente não encontrado"}), 404
        
        # Extrair informações
        processos_criados = data.get('processos_criados', 0)
        ids_criados = data.get('ids_criados', [])
        erros = data.get('erros', [])
        email_destino = cliente.get('email') or data.get('email')
        
        if not email_destino:
            return jsonify({
                "status": "aviso",
                "mensagem": "Cliente não possui email cadastrado"
            }), 200
        
        # Construir corpo do email
        html_body = f"""
        <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <h2>Relatório de Upload em Massa de Processos</h2>
                
                <p>Prezado/a,</p>
                
                <p>O upload em massa de processos foi concluído com sucesso!</p>
                
                <div style="background-color: #f5f5f5; padding: 15px; border-radius: 5px; margin: 15px 0;">
                    <p><strong>Data/Hora:</strong> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
                    <p><strong>Cliente:</strong> {cliente.get('nome_completo', 'N/A')}</p>
                    <p><strong>Processos Criados:</strong> <span style="color: green; font-weight: bold;">{processos_criados}</span></p>
                </div>
                
                <h3>Detalhes dos Processos Criados:</h3>
                <ul style="background-color: #f9f9f9; padding: 15px; border-left: 4px solid #4CAF50;">
        """
        
        for proc_id in ids_criados[:20]:  # Limitar a 20 para não sobrecarregar email
            html_body += f"<li><code>{proc_id}</code></li>"
        
        if len(ids_criados) > 20:
            html_body += f"<li><em>... e mais {len(ids_criados) - 20} processos</em></li>"
        
        html_body += """
                </ul>
        """
        
        if erros:
            html_body += f"""
                <h3 style="color: #d9534f;">Erros Encontrados:</h3>
                <ul style="background-color: #fff3cd; padding: 15px; border-left: 4px solid #ffc107;">
            """
            for erro in erros[:10]:  # Limitar a 10 erros
                html_body += f"<li>{erro}</li>"
            
            if len(erros) > 10:
                html_body += f"<li><em>... e mais {len(erros) - 10} erros</em></li>"
            
            html_body += "</ul>"
        
        html_body += """
                <p style="margin-top: 20px; color: #666; font-size: 12px;">
                    <em>Este é um email automático gerado pelo sistema Advocacia e IA.</em>
                </p>
            </body>
        </html>
        """
        
        # Configurações SMTP (usar variáveis de ambiente)
        smtp_host = os.getenv('SMTP_HOST', 'localhost')
        smtp_port = int(os.getenv('SMTP_PORT', '587'))
        smtp_user = os.getenv('SMTP_USER', '')
        smtp_pass = os.getenv('SMTP_PASSWORD', '')
        email_from = os.getenv('EMAIL_FROM', 'noreply@advocacia-ia.local')
        
        # Se não houver configuração SMTP, apenas registrar
        if not smtp_host or smtp_host == 'localhost':
            logger.info(f"Email seria enviado para {email_destino} (SMTP não configurado)")
            return jsonify({
                "status": "aviso",
                "mensagem": "SMTP não configurado. Notificação não foi enviada."
            }), 200
        
        try:
            # Enviar email
            msg = MIMEMultipart('alternative')
            msg['Subject'] = f'Upload em Massa de Processos - {processos_criados} processos criados'
            msg['From'] = email_from
            msg['To'] = email_destino
            
            msg.attach(MIMEText(html_body, 'html'))
            
            with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as server:
                if os.getenv('SMTP_USE_TLS', 'true').lower() == 'true':
                    server.starttls()
                if smtp_user and smtp_pass:
                    server.login(smtp_user, smtp_pass)
                server.send_message(msg)
            
            logger.info(f"Email de notificação enviado para {email_destino}")
            
            return jsonify({
                "status": "sucesso",
                "mensagem": f"Email enviado para {email_destino}"
            }), 200
            
        except smtplib.SMTPException as smtp_err:
            logger.warning(f"Erro ao enviar email SMTP: {smtp_err}")
            return jsonify({
                "status": "aviso",
                "mensagem": f"Email não pôde ser enviado: {str(smtp_err)}"
            }), 200
        except Exception as mail_err:
            logger.error(f"Erro geral ao enviar email: {mail_err}", exc_info=True)
            return jsonify({
                "status": "aviso",
                "mensagem": "Erro ao enviar notificação por email"
            }), 200
            
    except Exception as e:
        logger.error(f"Erro na notificação: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

# ============================================================
# FEATURE 3: Editar Processo com 12 Novos Campos (Item 1)
# ============================================================

@processos_bp.route('/<id_processo>/editar', methods=['GET'])
@login_required
def processo_edit_form(id_processo):
    """Formulário de edição do processo com 12 novos campos."""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        
        # Buscar processo
        processo = service.get_processo(id_processo)
        if not processo:
            flash('Processo não encontrado', 'danger')
            return redirect(url_for('clientes.list_clientes_ui'))
        
        # Buscar advogados para dropdown
        advogados = mgr.get_advogados()
        
        return render_template('processo_edit.html', processo=processo, advogados=advogados)
        
    except Exception as e:
        logger.error(f"Erro ao carregar formulário de edição: {e}", exc_info=True)
        flash(f'Erro ao carregar formulário: {e}', 'danger')
        return redirect(url_for('processos.painel_processo', id_processo=id_processo))

@processos_bp.route('/<id_processo>/salvar', methods=['POST'])
@login_required
def processo_save(id_processo):
    """Salva processo editado com 12 novos campos."""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        
        # Buscar processo atual para pegar id_cliente
        processo = service.get_processo(id_processo)
        if not processo:
            flash('Processo não encontrado', 'danger')
            return redirect(url_for('clientes.list_clientes_ui'))
        
        # Coletar dados do formulário
        dados = {
            'id_cliente': processo['id_cliente'],  # RealDictRow é dict, não objeto
            'nome_caso': request.form.get('nome_caso'),
            'numero_cnj': request.form.get('numero_cnj'),
            'status': request.form.get('status'),
            'advogado_oab': request.form.get('advogado_oab'),
            'tipo_parte': request.form.get('tipo_parte'),
            # 12 novos campos
            'local_tramite': request.form.get('local_tramite'),
            'comarca': request.form.get('comarca'),
            'area_atuacao': request.form.get('area_atuacao') or None,
            'instancia': request.form.get('instancia') or None,
            'subfase': request.form.get('subfase') or None,
            'assunto': request.form.get('assunto'),
            'valor_causa': request.form.get('valor_causa') or None,
            'data_distribuicao': request.form.get('data_distribuicao') or None,  # String vazia → None
            'data_encerramento': request.form.get('data_encerramento') or None,  # String vazia → None
            'sentenca': request.form.get('sentenca'),
            'em_execucao': request.form.get('em_execucao') == 'on',
            'segredo_justica': request.form.get('segredo_justica') == 'on'
        }
        
        # Salvar via CadastroManager (ordem: dados, id_processo)
        logger.info(f"[PROCESSO SAVE] Salvando processo {id_processo} com dados: {list(dados.keys())}")
        mgr.save_processo(dados, id_processo)
        logger.info(f"[PROCESSO SAVE] Processo {id_processo} salvo com sucesso!")
        
        flash('Processo atualizado com sucesso!', 'success')
        return redirect(url_for('processos.painel_processo', id_processo=id_processo))
        
    except Exception as e:
        logger.error(f"Erro ao salvar processo: {e}", exc_info=True)
        flash(f'Erro ao salvar processo: {e}', 'danger')
        return redirect(url_for('processos.processo_edit_form', id_processo=id_processo))

# ============================================================
# FEATURE 4: CRUD Partes Adversas (Item 3)
# ============================================================

@processos_bp.route('/<id_processo>/partes-adversas', methods=['GET'])
@login_required
def partes_adversas_page(id_processo):
    """Página de gerenciamento de partes adversas do processo."""
    try:
        # Buscar processo
        processo = service.get_processo(id_processo)
        if not processo:
            flash('Processo não encontrado', 'danger')
            return redirect(url_for('clientes.list_clientes_ui'))
        
        # Buscar partes adversas
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        partes = mgr.get_partes_adversas_by_processo(id_processo)
        
        return render_template('partes_adversas.html', processo=processo, partes=partes)
        
    except Exception as e:
        logger.error(f"Erro ao carregar partes adversas: {e}", exc_info=True)
        flash(f'Erro ao carregar partes adversas: {e}', 'danger')
        return redirect(url_for('processos.painel_processo', id_processo=id_processo))

@processos_bp.route('/api/<id_processo>/partes-adversas', methods=['GET'])
@login_required
def partes_adversas_api_list(id_processo):
    """API: Lista partes adversas do processo."""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        partes = mgr.get_partes_adversas_by_processo(id_processo)
        
        return jsonify({
            "status": "sucesso",
            "total": len(partes),
            "partes": [dict(p) for p in partes]
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao listar partes adversas: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

@processos_bp.route('/api/<id_processo>/partes-adversas', methods=['POST'])
@login_required
def partes_adversas_api_create(id_processo):
    """API: Cria nova parte adversa."""
    try:
        dados = request.get_json() or {}
        
        # Adicionar id_processo e tenant_id automaticamente
        dados['id_processo'] = id_processo
        
        tenant_id = getattr(g, 'tenant_id', None)
        dados['tenant_id'] = tenant_id  # CRÍTICO: tenant_id é NOT NULL
        
        logger.info(f"[PARTE ADVERSA CREATE] tenant_id={tenant_id}, dados={list(dados.keys())}")
        
        mgr = CadastroManager(tenant_id=tenant_id)
        
        id_parte = mgr.save_parte_adversa(dados)
        
        logger.info(f"Parte adversa criada: ID {id_parte} para processo {id_processo}")
        
        return jsonify({
            "status": "sucesso",
            "mensagem": "Parte adversa criada com sucesso",
            "id_parte": id_parte
        }), 201
        
    except ValueError as ve:
        return jsonify({"status": "erro", "mensagem": str(ve)}), 400
    except Exception as e:
        logger.error(f"Erro ao criar parte adversa: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

@processos_bp.route('/api/<id_processo>/partes-adversas/<int:id_parte>', methods=['GET'])
@login_required
def partes_adversas_api_get(id_processo, id_parte):
    """API: Busca parte adversa específica."""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        
        parte = mgr.get_parte_adversa_by_id(id_parte)
        
        if not parte:
            return jsonify({"status": "erro", "mensagem": "Parte adversa não encontrada"}), 404
        
        # Validar se parte pertence ao processo correto
        if parte.get('id_processo') != id_processo:
            return jsonify({"status": "erro", "mensagem": "Parte adversa não pertence a este processo"}), 403
        
        return jsonify({
            "status": "sucesso",
            "parte": dict(parte)
        }), 200
        
    except Exception as e:
        logger.error(f"Erro ao buscar parte adversa: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

@processos_bp.route('/api/<id_processo>/partes-adversas/<int:id_parte>', methods=['PUT'])
@login_required
def partes_adversas_api_update(id_processo, id_parte):
    """API: Atualiza parte adversa existente."""
    try:
        dados = request.get_json() or {}
        
        tenant_id = getattr(g, 'tenant_id', None)
        dados['tenant_id'] = tenant_id  # Garantir tenant_id no update também
        dados['id_processo'] = id_processo  # Garantir id_processo no update
        
        mgr = CadastroManager(tenant_id=tenant_id)
        
        # Validar se parte existe e pertence ao processo
        parte_atual = mgr.get_parte_adversa_by_id(id_parte)
        if not parte_atual:
            return jsonify({"status": "erro", "mensagem": "Parte adversa não encontrada"}), 404
        
        if parte_atual.get('id_processo') != id_processo:
            return jsonify({"status": "erro", "mensagem": "Parte adversa não pertence a este processo"}), 403
        
        # Atualizar
        mgr.save_parte_adversa(dados, id_parte=id_parte)
        
        logger.info(f"Parte adversa atualizada: ID {id_parte}")
        
        return jsonify({
            "status": "sucesso",
            "mensagem": "Parte adversa atualizada com sucesso"
        }), 200
        
    except ValueError as ve:
        return jsonify({"status": "erro", "mensagem": str(ve)}), 400
    except Exception as e:
        logger.error(f"Erro ao atualizar parte adversa: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

@processos_bp.route('/api/<id_processo>/partes-adversas/<int:id_parte>', methods=['DELETE'])
@login_required
def partes_adversas_api_delete(id_processo, id_parte):
    """API: Exclui parte adversa."""
    try:
        tenant_id = getattr(g, 'tenant_id', None)
        mgr = CadastroManager(tenant_id=tenant_id)
        
        # Validar se parte existe e pertence ao processo
        parte_atual = mgr.get_parte_adversa_by_id(id_parte)
        if not parte_atual:
            return jsonify({"status": "erro", "mensagem": "Parte adversa não encontrada"}), 404
        
        if parte_atual.get('id_processo') != id_processo:
            return jsonify({"status": "erro", "mensagem": "Parte adversa não pertence a este processo"}), 403
        
        # Excluir
        sucesso = mgr.delete_parte_adversa(id_parte)
        
        if sucesso:
            logger.info(f"Parte adversa excluída: ID {id_parte}")
            return jsonify({
                "status": "sucesso",
                "mensagem": "Parte adversa excluída com sucesso"
            }), 200
        else:
            return jsonify({"status": "erro", "mensagem": "Falha ao excluir parte adversa"}), 500
        
    except Exception as e:
        logger.error(f"Erro ao excluir parte adversa: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500

@processos_bp.route('/api/viacep/<cep>', methods=['GET'])
@login_required
def viacep_proxy(cep):
    """Proxy para consulta ViaCEP (evita CORS no frontend)."""
    try:
        import requests
        import re
        
        # Validar formato CEP
        cep_limpo = re.sub(r'\D', '', cep)
        if len(cep_limpo) != 8:
            return jsonify({"status": "erro", "mensagem": "CEP inválido"}), 400
        
        # Consultar ViaCEP
        response = requests.get(f"https://viacep.com.br/ws/{cep_limpo}/json/", timeout=5)
        response.raise_for_status()
        
        dados = response.json()
        
        if dados.get("erro"):
            return jsonify({"status": "erro", "mensagem": "CEP não encontrado"}), 404
        
        return jsonify({
            "status": "sucesso",
            "endereco": {
                "logradouro": dados.get("logradouro", ""),
                "bairro": dados.get("bairro", ""),
                "cidade": dados.get("localidade", ""),
                "estado": dados.get("uf", ""),
                "cep": dados.get("cep", "")
            }
        }), 200
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Erro ao consultar ViaCEP: {e}")
        return jsonify({"status": "erro", "mensagem": "Erro ao consultar CEP"}), 500
    except Exception as e:
        logger.error(f"Erro no proxy ViaCEP: {e}", exc_info=True)
        return jsonify({"status": "erro", "mensagem": str(e)}), 500
