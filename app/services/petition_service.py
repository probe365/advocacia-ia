# app/services/petition_service.py
import io
import json
import logging
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from psycopg2.errors import DatatypeMismatch, InvalidTextRepresentation, UniqueViolation
from docx import Document  # python-docx
from jinja2 import Template

import codecs

from cadastro_manager import CadastroManager

logger = logging.getLogger(__name__)

DEFAULT_CONTENT: Dict[str, Any] = {
    "texto_livre": "",
    "enderecamento": "",
    "tipo_acao": "",
    "fatos": "",
    "fundamentos": "",
    "pedidos": "",
    "provas": "",
    "local_data_assinatura": "",
}

# Templates “limpos” em Jinja (usa dados do processo/autor/reu/advogado)
TEMPLATES: Dict[str, str] = {
    "inicial": """EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {{ processo.local_tramite or "___ª VARA CÍVEL" }} DA COMARCA DE {{ processo.comarca or "____________" }}/__.

{% if processo.numero_cnj %}Processo nº {{ processo.numero_cnj }}{% endif %}

{{ autor.nome_completo or "NOME DO AUTOR" }}, {{ autor.nacionalidade or "nacionalidade" }}, {{ autor.estado_civil or "estado civil" }}, {{ autor.profissao or "profissão" }}, portador(a) do CPF/CNPJ nº {{ autor.cpf_cnpj or "__________" }}, residente e domiciliado(a) em {{ autor.endereco_completo or "ENDEREÇO" }}, e-mail {{ autor.email or "__________" }}, por intermédio de seu advogado(a) {{ advogado.nome or "NOME DO ADVOGADO" }} (OAB {{ advogado.oab or "____" }}), vem, respeitosamente, à presença de Vossa Excelência, propor a presente

AÇÃO DE: {{ processo.nome_caso or "__________" }}

em face de

{{ reu.nome_completo or "NOME DO RÉU" }}, CPF/CNPJ nº {{ reu.cpf_cnpj or "__________" }}, com endereço em {{ reu.endereco_completo or reu.endereco or "ENDEREÇO" }}, e-mail {{ reu.email or "__________" }}, pelos fatos e fundamentos a seguir expostos.

I. DOS FATOS
{{ content.fatos or "" }}

II. DO DIREITO
{{ content.fundamentos or "" }}

III. DOS PEDIDOS
{{ content.pedidos or "" }}

IV. DAS PROVAS
{{ content.provas or "Protesta provar o alegado por todos os meios de prova em direito admitidos." }}

V. DO VALOR DA CAUSA
{% if processo.valor_causa %}Dá-se à causa o valor de R$ {{ processo.valor_causa }}.{% else %}Dá-se à causa o valor de R$ __________.{% endif %}

Termos em que, pede deferimento.

{{ processo.comarca or "____________" }}, ____/____/______.

{{ advogado.nome or "________________________________" }}
OAB {{ advogado.oab or "____" }}
""",

    "contestacao": """EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {{ processo.local_tramite or "___ª VARA CÍVEL" }} DA COMARCA DE {{ processo.comarca or "____________" }}/__.

{% if processo.numero_cnj %}Processo nº {{ processo.numero_cnj }}{% endif %}

{{ reu.nome_completo or "NOME DO RÉU" }}, CPF/CNPJ nº {{ reu.cpf_cnpj or "__________" }}, com endereço em {{ reu.endereco_completo or reu.endereco or "ENDEREÇO" }}, por seu advogado(a) {{ advogado.nome or "NOME DO ADVOGADO" }} (OAB {{ advogado.oab or "____" }}), nos autos da ação movida por {{ autor.nome_completo or "NOME DO AUTOR" }}, vem apresentar

CONTESTAÇÃO

I. SÍNTESE DA DEMANDA
{{ content.fatos or "" }}

II. PRELIMINARES
{{ content.fundamentos or "" }}

III. MÉRITO
{{ content.pedidos or "" }}

IV. PROVAS
{{ content.provas or "Protesta provar o alegado por todos os meios de prova em direito admitidos." }}

Termos em que, pede deferimento.

{{ processo.comarca or "____________" }}, ____/____/______.

{{ advogado.nome or "________________________________" }}
OAB {{ advogado.oab or "____" }}
""",

    "replica": """EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {{ processo.local_tramite or "___ª VARA CÍVEL" }} DA COMARCA DE {{ processo.comarca or "____________" }}/__.

{% if processo.numero_cnj %}Processo nº {{ processo.numero_cnj }}{% endif %}

{{ autor.nome_completo or "NOME DO AUTOR" }}, já qualificado(a), nos autos em epígrafe, em face de {{ reu.nome_completo or "NOME DO RÉU" }}, vem apresentar

RÉPLICA À CONTESTAÇÃO

I. IMPUGNAÇÃO ÀS PRELIMINARES
{{ content.fundamentos or "" }}

II. IMPUGNAÇÃO AO MÉRITO
{{ content.fatos or "" }}

III. REQUERIMENTOS FINAIS
{{ content.pedidos or "" }}

Termos em que, pede deferimento.

{{ processo.comarca or "____________" }}, ____/____/______.

{{ advogado.nome or "________________________________" }}
OAB {{ advogado.oab or "____" }}
""",
}


class PetitionService:
    def __init__(self) -> None:
        pass

    # ----------------------------
    # Infra / helpers
    # ----------------------------
    def _require_tenant_id(self, tenant_id: Optional[str]) -> str:
        normalized = (str(tenant_id).strip() if tenant_id is not None else "")
        if not normalized or normalized.lower() == "none":
            normalized = os.getenv("DEFAULT_TENANT_ID", "public").strip()
        if not normalized:
            raise ValueError("tenant_id obrigatório")
        return normalized

    def _mgr(self, tenant_id: str) -> CadastroManager:
        return CadastroManager(tenant_id=self._require_tenant_id(tenant_id))

    def _norm_petition_type(self, petition_type: str) -> str:
        return (petition_type or "").strip().lower()

    def _ensure_dict_content(self, content: Any) -> Dict[str, Any]:
        if content is None:
            return {}
        if isinstance(content, dict):
            return content
        if isinstance(content, str):
            try:
                val = json.loads(content)
                return val if isinstance(val, dict) else {}
            except Exception:
                return {}
        return {}

    def _merge_content(self, existing: Optional[Dict[str, Any]], updates: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        merged = dict(DEFAULT_CONTENT)
        merged.update(existing or {})
        merged.update(updates or {})
        return merged

    def _empty_context(self) -> Dict[str, Any]:
        return {"processo": {}, "autor": {}, "reu": {}, "advogado": {}}

    def _get_context(self, tenant_id: str, process_id: Optional[str]) -> Dict[str, Any]:
        if not process_id:
            return self._empty_context()
        return self.preparar_contexto_peticao(tenant_id, process_id) or self._empty_context()

    def _process_id_candidates(self, process_id: Optional[str]) -> List[str]:
        if not process_id:
            return []
        pid = str(process_id).strip()
        if not pid:
            return []
        candidates = [pid]
        if pid.startswith("caso_"):
            candidates.append(pid.replace("caso_", "", 1))
        else:
            candidates.append(f"caso_{pid}")
        # unique, preserve order
        seen = set()
        out: List[str] = []
        for c in candidates:
            if c and c not in seen:
                seen.add(c)
                out.append(c)
        return out

    def _fmt_date_br(self, dt: Optional[datetime] = None) -> str:
        dt = dt or datetime.now()
        return dt.strftime("%d/%m/%Y")

    def _fmt_dt_br(self, dt: Optional[datetime] = None) -> str:
        dt = dt or datetime.now()
        return dt.strftime("%d/%m/%Y %H:%M")
    
    def _decode_unicode_escapes(self, s: str) -> str:
        if not s:
            return s
        # aplica até 2 passes (cobre \\u00e7 -> \u00e7 -> ç)
        out = s
        for _ in range(2):
            if "\\u" in out or "\\x" in out:
                try:
                    out = codecs.decode(out, "unicode_escape")
                except Exception:
                    break
            else:
                break
        return out

    def _firac_list_to_text(self, val: Any) -> str:
        if val is None:
            return ""

        # LIST real
        if isinstance(val, list):
            items = []
            for x in val:
                s = self._decode_unicode_escapes(str(x).strip())
                if s:
                    items.append(f"- {s}")
            return "\n".join(items)

        # STRING
        if isinstance(val, str):
            raw = val.strip()

            # JSON de lista?
            if raw.startswith("[") and raw.endswith("]"):
                try:
                    parsed = json.loads(raw)
                    if isinstance(parsed, list):
                        items = []
                        for x in parsed:
                            s = self._decode_unicode_escapes(str(x).strip())
                            if s:
                                items.append(f"- {s}")
                        return "\n".join(items)
                except Exception:
                    pass

            # texto simples
            return self._decode_unicode_escapes(raw)

        return str(val)

    def _summarize_firac_text(
        self,
        text: str,
        *,
        max_items: int = 5,
        lead_in: Optional[str] = None,
        tone: str = "formal",
    ) -> str:
        """
        Resume texto do FIRAC (lista ou parágrafos) em poucas ideias úteis,
        com tom formal/jurídico.
        """
        raw = (text or "").strip()
        if not raw:
            return ""

        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        items: List[str] = []
        for ln in lines:
            if ln.startswith("-"):
                ln = ln.lstrip("-").strip()
            if ln:
                items.append(ln)

        if not items:
            items = [raw]

        summary = self._select_strong_arguments(items, max_items)

        if tone == "formal":
            if len(summary) == 1:
                text_out = f"Em síntese, {summary[0].rstrip('.')} .".replace(" .", ".")
            else:
                bullets = "\n".join([f"- {s.rstrip('.')} .".replace(" .", ".") for s in summary])
                text_out = bullets
        else:
            text_out = summary[0] if len(summary) == 1 else " ".join(summary)

        if lead_in:
            return f"{lead_in}\n{text_out}"
        return text_out

    def _select_strong_arguments(self, items: List[str], max_items: int) -> List[str]:
        if not items:
            return []

        keywords = (
            "indevido",
            "abusivo",
            "ilegal",
            "nulo",
            "nulidade",
            "violação",
            "descumprimento",
            "inadimplemento",
            "prova",
            "documento",
            "notificação",
            "dano",
            "prejuízo",
            "cobrança",
            "responsabilidade",
            "direito",
            "dever",
            "contrato",
            "cláusula",
            "boa-fé",
            "fraude",
            "omissão",
            "recusa",
            "negativa",
        )

        def score(text: str) -> int:
            t = (text or "").lower()
            s = 0
            s += sum(2 for k in keywords if k in t)
            s += 1 if any(ch.isdigit() for ch in t) else 0
            s += 1 if len(t) >= 120 else 0
            return s

        ranked = sorted(items, key=score, reverse=True)
        return ranked[:max_items]

    # ----------------------------
    # Normalização de partes (para o template Jinja)
    # ----------------------------
    def _normalize_party_dict(self, p: Dict[str, Any]) -> Dict[str, Any]:
        """
        Garante que o dict do autor/réu tenha as chaves esperadas no template:
        nome_completo, cpf_cnpj, endereco_completo, email etc.
        """
        if not p:
            return {}

        # fontes possíveis: clientes, partes_adversas ou join de participants
        nome = (
                p.get("nome_completo")
                or p.get("nome")
                or p.get("nome_cliente")
                or p.get("razao_social")
                or p.get("razao")
                or p.get("nome_fantasia")
                or ""
            )

        cpf = (
            p.get("cpf_cnpj")
            or p.get("cpf")
            or p.get("cnpj")
            or p.get("documento")
            or ""
        )

        # se você tiver campos separados de endereço, monte um endereco_completo:
        end = p.get("endereco_completo") or p.get("endereco") or ""
        if not end:
            parts = []
            for k in ("logradouro", "numero", "complemento", "bairro", "cidade", "uf", "cep"):
                v = p.get(k)
                if v:
                    parts.append(str(v).strip())
            end = ", ".join(parts).strip()

        
        
        
        # nome = p.get("nome_completo") or p.get("nome") or p.get("razao_social") or p.get("nome_fantasia") or ""
        # cpf = p.get("cpf_cnpj") or p.get("cpf") or p.get("cnpj") or ""
        # end = p.get("endereco_completo") or p.get("endereco") or ""
        email = p.get("email") or ""

        out = dict(p)
        out["nome_completo"] = nome
        out["cpf_cnpj"] = cpf
        out["endereco_completo"] = end
        out["email"] = email
        return out

    # ----------------------------
    # FIRAC -> content (DOS FATOS / DO DIREITO / PEDIDOS)
    # ----------------------------
    def get_firac_by_process(self, tenant_id: str, process_id: str) -> Optional[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        candidates = self._process_id_candidates(process_id)
        if not candidates:
            return None
        row = mgr._execute_query(
            """
            SELECT facts, issue, rules, application, conclusion, source, updated_at
            FROM process_firac
            WHERE tenant_id = %s AND process_id = ANY(%s)
            LIMIT 1
            """,
            (str(tenant_id), candidates),
            fetch="one",
        )
        if not row:
            return None

        # Monta no formato esperado pelo apply_firac_autofill/_build_content_from_firac
        return {
            "facts": row.get("facts") or "",
            "issue": row.get("issue") or "",
            "rules": row.get("rules") or "",
            "application": row.get("application") or "",
            "conclusion": row.get("conclusion") or "",
            "source": row.get("source") or "",
        }
    def _build_content_from_firac(self, firac_data: Dict[str, Any], petition_type_norm: str) -> Dict[str, Any]:
        """
        Opção A: usa diretamente FIRAC (já gerado na tela/memória) e prepara texto.
        - inicial: fatos = FIRAC facts, fundamentos = rules+application, pedidos = conclusion (+ pedidos padrão)
        - contestacao/replica: usa FIRAC como base para síntese/impugnação (mínimo útil)
        """
        d = firac_data or {}
        facts = self._firac_list_to_text(d.get("facts"))
        rules = self._firac_list_to_text(d.get("rules"))
        application = (d.get("application") or "").strip()
        issue = (d.get("issue") or "").strip()
        conclusion = (d.get("conclusion") or "").strip()

        # blocos prontos (sem LLM aqui)
        fatos_txt = self._summarize_firac_text(
            facts,
            max_items=5,
            lead_in="Os fatos relevantes que ensejam a presente demanda podem ser assim sintetizados:",
            tone="formal",
        )
        fundamentos_txt = ""
        if issue:
            fundamentos_txt += f"Questão jurídica (síntese):\n{issue}\n\n"
        if rules:
            fundamentos_txt += (
                "Regras aplicáveis (resumo):\n"
                + self._summarize_firac_text(rules, max_items=4, tone="formal")
                + "\n\n"
            )
        if application:
            fundamentos_txt += (
                "Aplicação ao caso (síntese):\n"
                + self._summarize_firac_text(application, max_items=4, tone="formal")
                + "\n"
            )

        pedidos_txt = ""
        if conclusion:
            pedidos_txt += self._summarize_firac_text(
                conclusion,
                max_items=4,
                lead_in="Em conclusão, requer-se:",
                tone="formal",
            )
            pedidos_txt += "\n\n"

        # pedidos padrão (para não ficar vazio)
        pedidos_padrao = (
            "a) a citação da parte ré para, querendo, contestar a presente ação;\n"
            "b) a total procedência dos pedidos, nos termos da fundamentação;\n"
            "c) a condenação da parte ré ao pagamento de custas e honorários advocatícios;\n"
            "d) protesta provar o alegado por todos os meios de prova em direito admitidos.\n"
        )

        if petition_type_norm == "inicial":
            return {
                "fatos": fatos_txt,
                "fundamentos": fundamentos_txt,
                "pedidos": (pedidos_txt + pedidos_padrao).strip(),
            }

        if petition_type_norm == "contestacao":
            # contestação: fatos = síntese, fundamentos = preliminares, pedidos = mérito
            return {
                "fatos": fatos_txt,
                "fundamentos": fundamentos_txt,
                "pedidos": (pedidos_txt or "Requer-se a improcedência dos pedidos, com condenação em custas e honorários.").strip(),
            }

        # réplica
        return {
            "fatos": fatos_txt,
            "fundamentos": fundamentos_txt,
            "pedidos": (pedidos_txt or "Reitera-se a procedência dos pedidos iniciais, rechaçando-se as alegações defensivas.").strip(),
        }

    # ----------------------------
    # Contexto: processo/cliente/réu/advogado (via participants)
    # ----------------------------
    def rerender_if_placeholders(self, tenant_id: str, petition_id: int) -> None:
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            return
        content = petition.get("content") or {}
        texto = (content.get("texto_livre") or "")

        # sinais clássicos de placeholder
        if "NOME DO AUTOR" not in texto and "ENDEREÇO" not in texto:
            return

        ptype = self._norm_petition_type(petition.get("petition_type") or "inicial")
        ctx = self._get_context(tenant_id, petition.get("process_id"))
        rendered = self._render_template_text(ptype, ctx, content)
        self.save_petition_content(tenant_id, petition_id, {"texto_livre": rendered})

    
    
    def preparar_contexto_peticao(self, tenant_id: str, process_id: str) -> Optional[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)

        proc = mgr.get_processo_by_id(process_id)
        if not proc:
            return None

        advogado = mgr.get_advogado_by_oab(proc.get("advogado_oab")) if proc.get("advogado_oab") else None
        cliente = mgr.get_cliente_by_id(proc.get("id_cliente")) if proc.get("id_cliente") else None

        participants = mgr.get_process_participants(process_id) or []

        def pick_primary(role: str) -> Optional[Dict[str, Any]]:
            role_norm = (role or "").lower()
            aliases = {
                "autor": ["autor", "reclamante"],
                "reu": ["reu", "reclamada"],
            }.get(role_norm, [role_norm])

            p = next(
                (
                    x
                    for x in participants
                    if str(x.get("role") or "").lower() in aliases and x.get("is_primary")
                ),
                None,
            )
            if p:
                return p
            return next(
                (x for x in participants if str(x.get("role") or "").lower() in aliases),
                None,
            )

        # papel do cliente no processo determina o “lado”
        client_role = (proc.get("tipo_parte") or "autor").lower()
        opposite_role = "reu" if client_role == "autor" else ("autor" if client_role == "reu" else "reu")

        autor_p = pick_primary("autor")
        reu_p = pick_primary("reu")

        # resolve participante -> tabela real
        autor = self._resolve_party(mgr, autor_p) if autor_p else {}
        reu = self._resolve_party(mgr, reu_p) if reu_p else {}

        # fallback: se cliente é autor/reu e não veio via participants (por seed incompleto)
        # fallback forte: se existir cliente no processo, ele SEMPRE representa o polo do cliente
        if cliente:
            if client_role == "autor":
                autor = autor or cliente
            elif client_role == "reu":
                reu = reu or cliente
            else:
                # se não definido, assume autor
                autor = autor or cliente


        # fallback para polo oposto via participants
        if client_role == "autor" and not reu:
            opp = pick_primary(opposite_role)
            reu = self._resolve_party(mgr, opp) if opp else {}
        if client_role == "reu" and not autor:
            opp = pick_primary(opposite_role)
            autor = self._resolve_party(mgr, opp) if opp else {}

        # fallback final: usar partes_adversas diretamente (quando participants não existem)
        if client_role == "autor" and not reu:
            reu = self._pick_adverso_for_role(mgr, process_id, "reu")
        if client_role == "reu" and not autor:
            autor = self._pick_adverso_for_role(mgr, process_id, "autor")

        autor = self._normalize_party_dict(autor or {})
        reu = self._normalize_party_dict(reu or {})
        advogado = advogado or {}
        proc = proc or {}

        return {
            "processo": proc,
            "autor": autor,
            "reu": reu,
            "advogado": advogado,
            "participants": participants,
        }

    def _pick_adverso_for_role(
        self,
        mgr: CadastroManager,
        process_id: str,
        role: str,
    ) -> Dict[str, Any]:
        adversos = mgr.get_partes_adversas_by_processo(process_id) or []
        if not adversos:
            return {}

        role_norm = (role or "").lower()
        match = next(
            (p for p in adversos if str(p.get("tipo_parte") or "").lower() == role_norm),
            None,
        )
        return match or adversos[0]

    def _resolve_party(self, mgr: CadastroManager, participant: Dict[str, Any]) -> Dict[str, Any]:
        if not participant:
            return {}
        kind = (participant.get("party_kind") or "").lower()
        pid = participant.get("party_id")

        if kind == "cliente":
            try:
                return mgr.get_cliente_by_id(int(pid)) or {}
            except Exception:
                return mgr.get_cliente_by_id(pid) or {}


        if kind == "adverso":
            row = mgr._execute_query(
                """
                SELECT *
                FROM partes_adversas
                WHERE tenant_id=%s AND id=%s
                LIMIT 1
                """,
                (mgr.tenant_id, int(pid)),
                fetch="one",
            )
            return row or {}

        return {}

    # ----------------------------
    # Renderização do modelo Jinja
    # ----------------------------
    def _render_template_text(self, petition_type_norm: str, ctx: Dict[str, Any], content: Dict[str, Any]) -> str:
        tpl = TEMPLATES.get(petition_type_norm)
        if not tpl:
            raise ValueError(f"Template não encontrado para petition_type={petition_type_norm}")
        ctx2 = dict(ctx or {})
        merged = dict(DEFAULT_CONTENT)
        merged.update(content or {})
        ctx2["content"] = merged
        return Template(tpl).render(**ctx2)

    # ----------------------------
    # Create-or-get (evita duplicata)
    # ----------------------------
    def create_or_get_petition(self, tenant_id: str, user_id: int, process_id: str, petition_type: str) -> int:
        mgr = self._mgr(tenant_id)
        ptype = self._norm_petition_type(petition_type)

        row = mgr._execute_query(
            """
            SELECT id
              FROM petitions
             WHERE tenant_id = %s
               AND process_id = %s
               AND LOWER(petition_type) = %s
             LIMIT 1
            """,
            (str(tenant_id), str(process_id), ptype),
            fetch="one",
        )

        if row and row.get("id") is not None:
            return int(row["id"])

        try:
            self.create_petition(
                tenant_id=str(tenant_id),
                user_id=int(user_id),
                process_id=str(process_id),
                petition_type=ptype,
            )
        except UniqueViolation:
            logger.info(
                "petition created concurrently; will re-fetch",
                extra={"tenant_id": tenant_id, "process_id": process_id, "ptype": ptype},
            )

        row2 = mgr._execute_query(
            """
            SELECT id
              FROM petitions
             WHERE tenant_id = %s
               AND process_id = %s
               AND LOWER(petition_type) = %s
             ORDER BY created_at DESC, id DESC
             LIMIT 1
            """,
            (str(tenant_id), str(process_id), ptype),
            fetch="one",
        )
        if not row2 or row2.get("id") is None:
            raise RuntimeError("Falha ao criar petição (ID não retornado).")

        return int(row2["id"])

    # ----------------------------
    # CRUD petitions
    # ----------------------------
    def create_petition(self, tenant_id: str, user_id: int, process_id: str, petition_type: str) -> None:
        mgr = self._mgr(tenant_id)

        proc = mgr.get_processo_by_id(process_id)
        raw_cliente_id = proc.get("id_cliente") if proc else None
        cliente_id = str(raw_cliente_id) if raw_cliente_id not in (None, "") else None

        content = dict(DEFAULT_CONTENT)
        ptype = self._norm_petition_type(petition_type)

        # cria um texto inicial mínimo usando o template Jinja (ainda sem FIRAC)
        ctx = self._get_context(tenant_id, process_id)
        logger.warning("CTX PETICAO", extra={
            "process_id": process_id,
            "tenant_id": tenant_id,
            "autor": ctx.get("autor"),
            "reu": ctx.get("reu"),
            "advogado": ctx.get("advogado"),
            "processo_id_cliente": (ctx.get("processo") or {}).get("id_cliente"),
            "tipo_parte": (ctx.get("processo") or {}).get("tipo_parte"),
        })
        content["texto_livre"] = self._render_template_text(ptype, ctx, content)

        try:
            self._insert_petition_row(
                mgr,
                tenant_id=tenant_id,
                user_id=user_id,
                cliente_id=cliente_id,
                process_id=process_id,
                petition_type=ptype,
                content=content,
            )
        except (InvalidTextRepresentation, DatatypeMismatch):
            logger.warning(
                "cliente_id mismatch; inserting with NULL cliente_id",
                extra={"tenant_id": tenant_id, "process_id": process_id, "ptype": ptype},
            )
            self._insert_petition_row(
                mgr,
                tenant_id=tenant_id,
                user_id=user_id,
                cliente_id=None,
                process_id=process_id,
                petition_type=ptype,
                content=content,
            )

    def _insert_petition_row(
        self,
        mgr: CadastroManager,
        *,
        tenant_id: str,
        user_id: int,
        cliente_id: Optional[str],
        process_id: str,
        petition_type: str,
        content: Dict[str, Any],
    ) -> None:
        mgr._execute_query(
            """
            INSERT INTO petitions (tenant_id, user_id, cliente_id, process_id, petition_type, content, status)
            VALUES (%s, %s, %s, %s, %s, %s::jsonb, 'draft')
            """,
            (str(tenant_id), int(user_id), cliente_id, str(process_id), petition_type, json.dumps(content, ensure_ascii=False), ),
            fetch=None,
        )

    def list_by_process(self, tenant_id: str, process_id: str) -> List[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        rows = mgr._execute_query(
            """
            SELECT id, tenant_id, user_id, cliente_id, process_id, petition_type, content, status, created_at, updated_at
              FROM petitions
             WHERE tenant_id = %s AND process_id = %s
             ORDER BY updated_at DESC NULLS LAST, id DESC
            """,
            (str(tenant_id), str(process_id)),
            fetch="all",
        ) or []

        for r in rows:
            r["content"] = self._merge_content(self._ensure_dict_content(r.get("content")), None)

        return rows

    def get_petition(self, tenant_id: str, petition_id: int) -> Optional[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        row = mgr._execute_query(
            """
            SELECT id, tenant_id, user_id, cliente_id, process_id, petition_type, content, status, created_at, updated_at
              FROM petitions
             WHERE tenant_id = %s AND id = %s
             LIMIT 1
            """,
            (str(tenant_id), int(petition_id)),
            fetch="one",
        )
        if not row:
            return None

        row["content"] = self._merge_content(self._ensure_dict_content(row.get("content")), None)
        return row

    def save_petition_content(self, tenant_id: str, petition_id: int, content_updates: Dict[str, Any]) -> None:
        mgr = self._mgr(tenant_id)
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            raise ValueError("Petição não encontrada")

        content = self._merge_content(petition.get("content") or {}, content_updates)

        mgr._execute_query(
            """
            UPDATE petitions
               SET content = %s::jsonb,
                   updated_at = CURRENT_TIMESTAMP
             WHERE tenant_id = %s AND id = %s
            """,
            (json.dumps(content, ensure_ascii=False), str(tenant_id), int(petition_id)),
            fetch=None,
        )

    def set_status(self, tenant_id: str, petition_id: int, status: str) -> None:
        mgr = self._mgr(tenant_id)
        status_norm = (status or "").strip().lower()
        if status_norm not in ("draft", "final"):
            status_norm = "draft"

        mgr._execute_query(
            """
            UPDATE petitions
               SET status = %s,
                   updated_at = CURRENT_TIMESTAMP
             WHERE tenant_id = %s AND id = %s
            """,
            (status_norm, str(tenant_id), int(petition_id)),
            fetch=None,
        )

    def finalize_petition(self, tenant_id: str, petition_id: int) -> None:
        self.set_status(tenant_id, petition_id, "final")

    # ----------------------------
    # Aplicar FIRAC na petição (Opção A)
    # ----------------------------
    def apply_firac_autofill(self, tenant_id: str, petition_id: int, firac_data: Dict[str, Any], *, force: bool = True) -> None:
        """
        Preenche content.fatos/fundamentos/pedidos a partir do FIRAC e re-renderiza texto_livre.
        - force=True: sobrescreve os campos e o texto_livre.
        - force=False: só completa o que estiver vazio.
        """
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            raise ValueError("Petição não encontrada")

        ptype = self._norm_petition_type(petition.get("petition_type") or "inicial")
        current = petition.get("content") or {}

        built = self._build_content_from_firac(firac_data or {}, ptype)

        updates: Dict[str, Any] = {}
        for k in ("fatos", "fundamentos", "pedidos"):
            if force or not (current.get(k) or "").strip():
                updates[k] = built.get(k, "")

        # re-render texto_livre com template jinja (agora completo)
        ctx = self._get_context(tenant_id, petition.get("process_id"))
        merged = self._merge_content(current, updates)
        texto = self._render_template_text(ptype, ctx, merged)
        updates["texto_livre"] = texto

        self.save_petition_content(tenant_id, petition_id, updates)

    # ----------------------------
    # Aplicar template se vazio
    # ----------------------------
    def apply_model_text_if_empty(self, tenant_id: str, petition_id: int) -> None:
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            return

        content = petition.get("content") or {}
        texto = (content.get("texto_livre") or "").strip()
        if texto:
            return

        ptype = self._norm_petition_type(petition.get("petition_type") or "")
        ctx = self._get_context(tenant_id, petition.get("process_id"))
        rendered = self._render_template_text(ptype, ctx, content)

        self.save_petition_content(
            tenant_id=tenant_id,
            petition_id=petition_id,
            content_updates={"texto_livre": rendered},
        )

    def autofill_petition_from_firac_if_possible(
        self,
        tenant_id: str,
        petition_id: int,
        *,
        force: bool = False,
    ) -> None:
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            return

        process_id = petition.get("process_id")
        if not process_id:
            return

        firac = self.get_firac_by_process(tenant_id, str(process_id))
        if not firac:
            return

        self.apply_firac_autofill(
            tenant_id=tenant_id,
            petition_id=petition_id,
            firac_data=firac,
            force=force,
        )

    def autofill_petition_text(self, tenant_id: str, petition_id: int) -> None:
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            return

        content = petition.get("content") or {}
        ptype = self._norm_petition_type(petition.get("petition_type") or "inicial")
        ctx = self._get_context(tenant_id, petition.get("process_id"))
        rendered = self._render_template_text(ptype, ctx, content)
        self.save_petition_content(
            tenant_id=tenant_id,
            petition_id=petition_id,
            content_updates={"texto_livre": rendered},
        )

    # ----------------------------
    # Versionamento (mantido)
    # ----------------------------
    def list_versions(self, tenant_id: str, petition_id: int) -> List[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        rows = mgr._execute_query(
            """
            SELECT id, petition_id, version, rendered_content, created_at, created_by_user_id
              FROM petition_versions
             WHERE tenant_id = %s AND petition_id = %s
             ORDER BY version DESC, id DESC
            """,
            (str(tenant_id), int(petition_id)),
            fetch="all",
        ) or []
        return rows

    # ----------------------------
    # DOCX export
    # ----------------------------
    def build_docx_bytes(self, tenant_id: str, petition_id: int) -> bytes:
        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            raise ValueError("Petição não encontrada")

        process_id = petition.get("process_id")
        ctx = self._get_context(tenant_id, process_id)

        content = petition.get("content") or {}
        texto = (content.get("texto_livre") or "").strip()

        # fallback: se vazio, renderiza modelo
        if not texto:
            ptype = self._norm_petition_type(petition.get("petition_type") or "inicial")
            texto = self._render_template_text(ptype, ctx, content)

        doc = Document()

        titulo = f"{(petition.get('petition_type') or '').upper()} • Processo {ctx.get('processo', {}).get('numero_cnj') or ctx.get('processo', {}).get('id_processo') or ''}".strip(" •")
        doc.add_heading(titulo or "Petição", level=1)
        doc.add_paragraph(f"Criada/Atualizada: {self._fmt_dt_br()}")
        doc.add_paragraph("")

        for line in texto.splitlines():
            doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
