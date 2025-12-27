from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict

from docx import Document

logger = logging.getLogger(__name__)


class ExportService:
    @staticmethod
    def _export_dir(id_processo: str, tenant_id: str | None) -> Path:
        tenant_segment = str(tenant_id) if tenant_id else "default"
        export_dir = Path("./cases") / tenant_segment / str(id_processo) / "exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        return export_dir

    @staticmethod
    def to_docx(dados_peticao: Dict[str, Any], filename: str, id_processo: str, tenant_id: str | None) -> bool:
        """Gera um arquivo DOCX da petição e salva no diretório de exports do caso."""
        try:
            export_dir = ExportService._export_dir(id_processo=id_processo, tenant_id=tenant_id)
            path = export_dir / filename

            texto = (
                (dados_peticao.get("conteudo_peticao") or "").strip()
                or (dados_peticao.get("peticao") or "").strip()
                or (dados_peticao.get("narrativa_dos_fatos") or "").strip()
            )
            if not texto:
                texto = "Texto da petição não gerado."

            doc = Document()

            # Mantém quebras de linha: cada linha vira um parágrafo.
            for line in str(texto).splitlines():
                if line.strip() == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(line)

            doc.save(str(path))
            return True
        except Exception:
            logger.exception("Erro crítico ao gerar DOCX", extra={"id_processo": id_processo, "tenant_id": tenant_id})
            return False

    @staticmethod
    def generate_pdf(conteudo: str, titulo: str, path: Path) -> bool:
        """Gera PDF simples via FPDF (sem HTML/CSS) e salva em 'path'."""
        try:
            from fpdf import FPDF

            pdf = FPDF(format="A4")
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            pdf.set_font("Arial", "B", 14)
            pdf.multi_cell(0, 8, (titulo or "Documento").encode("latin-1", "replace").decode("latin-1"))
            pdf.ln(3)

            pdf.set_font("Arial", size=11)
            text = conteudo or ""
            for line in str(text).splitlines():
                safe_line = line.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 6, safe_line)

            path.parent.mkdir(parents=True, exist_ok=True)
            pdf.output(str(path))
            return True
        except Exception:
            logger.exception("Erro ao gerar PDF", extra={"path": str(path)})
            return False

