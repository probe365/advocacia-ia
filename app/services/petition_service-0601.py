# app/services/petition_service.py
import io
import json
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from psycopg2.errors import DatatypeMismatch, InvalidTextRepresentation, UniqueViolation
from jinja2 import Environment, BaseLoader

from cadastro_manager import CadastroManager

logger = logging.getLogger(__name__)


DEFAULT_CONTENT: Dict[str, Any] = {
    "texto_livre": "",          # campo livre (editor principal)
    "enderecamento": "",
    "tipo_acao": "",
    "fatos": "",
    "fundamentos": "",
    "pedidos": "",
    "provas": "",
    "local_data_assinatura": "",
}


def _safe_str(x: Any) -> str:
    if x is None:
        return ""
    return str(x)


class PetitionService:
    def __init__(self) -> None:
        # Jinja “simples” (sem autoescape, pois é texto jurídico)
        self._jinja = Environment(
            loader=BaseLoader(),
            autoescape=False,
            trim_blocks=True,
            lstrip_blocks=True,
        )

    # -----------------------------
    # Infra / helpers
    # -----------------------------
    def _mgr(self, tenant_id: str) -> CadastroManager:
        return CadastroManager(tenant_id=str(tenant_id))

    def _norm_petition_type(self, petition_type: str) -> str:
        # Normaliza para evitar duplicatas: INICIAL, Inicial, inicial etc.
        return (petition_type or "").strip().lower()

    def _ensure_dict_content(self, content: Any) -> Dict[str, Any]:
        if content is None:
            return {}
        if isinstance(content, dict):
            return content
        if isinstance(content, str):
            try:
                val = json.loads(content)
                return val if isinstance(val, dict) else {}
            except Exception:
                return {}
        return {}

    def _merge_content(self, content: Any) -> Dict[str, Any]:
        d = self._ensure_dict_content(content)
        merged = dict(DEFAULT_CONTENT)
        merged.update(d or {})
        return merged

    # -----------------------------
    # Contexto (processo/cliente/advogado/partes)
    # -----------------------------
    def preparar_contexto_peticao(self, id_processo: str, tenant_id: str) -> Optional[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)

        proc = mgr.get_processo_by_id(id_processo)
        if not proc:
            return None

        cliente = mgr.get_cliente_by_id(proc.get("id_cliente"))
        advogado = mgr.get_advogado_by_oab(proc.get("advogado_oab"))
        partes = mgr.get_partes_adversas_by_processo(id_processo) or []
        reu = partes[0] if partes else {}

        # “Extras” úteis que normalmente aparecem no texto
        numero_cnj = proc.get("numero_cnj") or proc.get("nro_cnj") or proc.get("cnj") or ""
        comarca = proc.get("comarca") or ""
        instancia = proc.get("instancia") or ""
        vara = proc.get("vara") or proc.get("local_tramite") or ""

        return {
            "processo": proc or {},
            "autor": cliente or {},
            "advogado": advogado or {},
            "reu": reu or {},
            "adversos": partes,
            "meta": {
                "numero_cnj": numero_cnj,
                "comarca": comarca,
                "vara": vara,
                "instancia": instancia,
            },
        }

    # -----------------------------
    # Create or get (evita duplicatas por tipo)
    # -----------------------------
    def create_or_get_petition(self, tenant_id: str, user_id: int, process_id: str, petition_type: str) -> int:
        mgr = self._mgr(tenant_id)
        ptype = self._norm_petition_type(petition_type)

        row = mgr._execute_query(
            """
            SELECT id
            FROM petitions
            WHERE tenant_id = %s
              AND process_id = %s
              AND LOWER(petition_type) = %s
            LIMIT 1
            """,
            (str(tenant_id), str(process_id), ptype),
            fetch="one",
        )

        if row and row.get("id") is not None:
            return int(row["id"])

        try:
            self.create_petition(
                tenant_id=str(tenant_id),
                user_id=int(user_id),
                process_id=str(process_id),
                petition_type=ptype,
            )
        except UniqueViolation:
            logger.info(
                "petition already created concurrently",
                extra={"tenant_id": tenant_id, "process_id": process_id, "petition_type": ptype},
            )

        row2 = mgr._execute_query(
            """
            SELECT id
            FROM petitions
            WHERE tenant_id = %s
              AND process_id = %s
              AND LOWER(petition_type) = %s
            ORDER BY created_at DESC, id DESC
            LIMIT 1
            """,
            (str(tenant_id), str(process_id), ptype),
            fetch="one",
        )

        if not row2 or row2.get("id") is None:
            raise RuntimeError("Falha ao criar petição (ID não retornado).")

        return int(row2["id"])

    # -----------------------------
    # CRUD petitions
    # -----------------------------
    def create_petition(self, tenant_id: str, user_id: int, process_id: str, petition_type: str) -> None:
        mgr = self._mgr(tenant_id)

        proc = mgr.get_processo_by_id(process_id)
        raw_cliente_id = proc.get("id_cliente") if proc else None
        cliente_id = str(raw_cliente_id) if raw_cliente_id not in (None, "") else None

        content = dict(DEFAULT_CONTENT)
        ptype = self._norm_petition_type(petition_type)

        try:
            self._insert_petition_row(
                mgr,
                tenant_id=tenant_id,
                user_id=user_id,
                cliente_id=cliente_id,
                process_id=process_id,
                petition_type=ptype,
                content=content,
            )
        except (InvalidTextRepresentation, DatatypeMismatch):
            logger.warning(
                "cliente_id fallback to NULL due to datatype mismatch",
                extra={"tenant_id": tenant_id, "process_id": process_id, "petition_type": ptype, "cliente_id": cliente_id},
            )
            self._insert_petition_row(
                mgr,
                tenant_id=tenant_id,
                user_id=user_id,
                cliente_id=None,
                process_id=process_id,
                petition_type=ptype,
                content=content,
            )

    def _insert_petition_row(
        self,
        mgr: CadastroManager,
        *,
        tenant_id: str,
        user_id: int,
        cliente_id: Optional[str],
        process_id: str,
        petition_type: str,
        content: Dict[str, Any],
    ) -> None:
        mgr._execute_query(
            """
            INSERT INTO petitions (tenant_id, user_id, cliente_id, process_id, petition_type, content, status)
            VALUES (%s, %s, %s, %s, %s, %s::jsonb, 'draft')
            """,
            (
                str(tenant_id),
                int(user_id),
                cliente_id,
                str(process_id),
                petition_type,
                json.dumps(content),
            ),
            fetch=None,
        )

    def list_by_process(self, tenant_id: str, process_id: str) -> List[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        rows = mgr._execute_query(
            """
            SELECT id, tenant_id, user_id, cliente_id, process_id, petition_type, content, status, created_at, updated_at
            FROM petitions
            WHERE tenant_id = %s AND process_id = %s
            ORDER BY updated_at DESC NULLS LAST, id DESC
            """,
            (str(tenant_id), str(process_id)),
            fetch="all",
        ) or []

        for r in rows:
            r["content"] = self._merge_content(r.get("content"))

        return rows

    def get_petition(self, tenant_id: str, petition_id: int) -> Optional[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        row = mgr._execute_query(
            """
            SELECT id, tenant_id, user_id, cliente_id, process_id, petition_type, content, status, created_at, updated_at
            FROM petitions
            WHERE tenant_id = %s AND id = %s
            LIMIT 1
            """,
            (str(tenant_id), int(petition_id)),
            fetch="one",
        )
        if not row:
            return None

        row["content"] = self._merge_content(row.get("content"))
        return row

    def save_petition_content(self, tenant_id: str, petition_id: int, content_updates: Dict[str, Any]) -> None:
        mgr = self._mgr(tenant_id)

        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            raise ValueError("Petição não encontrada")

        content = dict(DEFAULT_CONTENT)
        content.update(petition.get("content") or {})
        content.update(content_updates or {})

        mgr._execute_query(
            """
            UPDATE petitions
            SET content = %s::jsonb,
                updated_at = CURRENT_TIMESTAMP
            WHERE tenant_id = %s AND id = %s
            """,
            (json.dumps(content), str(tenant_id), int(petition_id)),
            fetch=None,
        )

    # -----------------------------
    # Versionamento
    # -----------------------------
    def list_versions(self, tenant_id: str, petition_id: int) -> List[Dict[str, Any]]:
        mgr = self._mgr(tenant_id)
        return mgr._execute_query(
            """
            SELECT id, petition_id, version, rendered_content, created_at, created_by_user_id
            FROM petition_versions
            WHERE tenant_id = %s AND petition_id = %s
            ORDER BY version DESC, id DESC
            """,
            (str(tenant_id), int(petition_id)),
            fetch="all",
        ) or []

    def _next_version_number(self, mgr: CadastroManager, tenant_id: str, petition_id: int) -> int:
        row = mgr._execute_query(
            """
            SELECT COALESCE(MAX(version), 0) AS maxv
            FROM petition_versions
            WHERE tenant_id = %s AND petition_id = %s
            """,
            (str(tenant_id), int(petition_id)),
            fetch="one",
        ) or {}
        return int(row.get("maxv") or 0) + 1

    def finalize_petition(self, tenant_id: str, petition_id: int) -> None:
        mgr = self._mgr(tenant_id)
        mgr._execute_query(
            """
            UPDATE petitions
            SET status = 'final',
                updated_at = CURRENT_TIMESTAMP
            WHERE tenant_id = %s AND id = %s
            """,
            (str(tenant_id), int(petition_id)),
            fetch=None,
        )

    # -----------------------------
    # Templates (3 modelos) + render
    # -----------------------------
    def _builtin_template_for(self, petition_type_norm: str) -> str:
        """
        🔥 Aqui é o lugar para você colar os 3 textos-modelo completos.
        Eu deixei um modelo base bem “jurídico”, mas você pode substituir por
        seus textos exatamente como você já tem.
        """
        if petition_type_norm == "inicial":
            return (
                "{{ content.enderecamento }}\n\n"
                "Processo nº: {{ meta.numero_cnj }}\n"
                "Comarca: {{ meta.comarca }}  Vara: {{ meta.vara }}  Instância: {{ meta.instancia }}\n\n"
                "{{ autor.nome_completo or autor.razao_social }}\n"
                "em face de {{ reu.nome_completo or reu.razao_social }}\n\n"
                "AÇÃO: {{ content.tipo_acao }}\n\n"
                "I - DOS FATOS\n{{ content.fatos }}\n\n"
                "II - DOS FUNDAMENTOS\n{{ content.fundamentos }}\n\n"
                "III - DOS PEDIDOS\n{{ content.pedidos }}\n\n"
                "IV - DAS PROVAS\n{{ content.provas }}\n\n"
                "{{ content.local_data_assinatura }}\n\n"
                "{{ advogado.nome }} - OAB {{ advogado.oab }}\n"
            )

        if petition_type_norm == "contestacao":
            return (
                "{{ content.enderecamento }}\n\n"
                "Processo nº: {{ meta.numero_cnj }}\n\n"
                "{{ reu.nome_completo or reu.razao_social }}, já qualificado(a), por seu advogado,\n"
                "apresenta CONTESTAÇÃO à ação movida por {{ autor.nome_completo or autor.razao_social }}.\n\n"
                "I - SÍNTESE\n{{ content.fatos }}\n\n"
                "II - PRELIMINARES / MÉRITO\n{{ content.fundamentos }}\n\n"
                "III - PEDIDOS\n{{ content.pedidos }}\n\n"
                "{{ content.local_data_assinatura }}\n\n"
                "{{ advogado.nome }} - OAB {{ advogado.oab }}\n"
            )

        # replica (default)
        return (
            "{{ content.enderecamento }}\n\n"
            "Processo nº: {{ meta.numero_cnj }}\n\n"
            "{{ autor.nome_completo or autor.razao_social }}, por seu advogado, apresenta RÉPLICA.\n\n"
            "I - DOS PONTOS CONTROVERTIDOS\n{{ content.fatos }}\n\n"
            "II - DO MÉRITO\n{{ content.fundamentos }}\n\n"
            "III - DOS PEDIDOS\n{{ content.pedidos }}\n\n"
            "{{ content.local_data_assinatura }}\n\n"
            "{{ advogado.nome }} - OAB {{ advogado.oab }}\n"
        )

    def render_from_template(self, tenant_id: str, user_id: int, petition_id: int) -> None:
        """
        Render determinístico (sem IA) usando:
        - contexto do processo/cliente/advogado/partes adversas
        - content (campos do editor)
        - modelo por tipo (inicial/contestacao/replica)
        Gera nova versão em petition_versions e NÃO sobrescreve texto_livre automaticamente,
        a menos que você queira (ver comentário abaixo).
        """
        mgr = self._mgr(tenant_id)

        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            raise ValueError("Petição não encontrada")

        ptype = self._norm_petition_type(petition.get("petition_type") or "")
        process_id = petition.get("process_id")
        if not process_id:
            raise ValueError("process_id ausente na petição")

        ctx = self.preparar_contexto_peticao(str(process_id), tenant_id)
        if not ctx:
            raise ValueError("Contexto do processo não encontrado")

        content = self._merge_content(petition.get("content"))

        # (Opcional) se texto_livre estiver vazio e campos estruturados também,
        # você pode gerar a partir do modelo mesmo assim.
        template_text = self._builtin_template_for(ptype)

        tpl = self._jinja.from_string(template_text)
        rendered = tpl.render(**ctx, content=content).strip()

        # versionamento
        nextv = self._next_version_number(mgr, tenant_id, petition_id)
        mgr._execute_query(
            """
            INSERT INTO petition_versions (tenant_id, petition_id, version, rendered_content, created_by_user_id)
            VALUES (%s, %s, %s, %s, %s)
            """,
            (str(tenant_id), int(petition_id), int(nextv), rendered, int(user_id)),
            fetch=None,
        )

        # ✅ Se você quiser que o texto gerado “caia” no editor automaticamente,
        # descomente este trecho:
        #
        # content2 = dict(content)
        # content2["texto_livre"] = rendered
        # mgr._execute_query(
        #     """
        #     UPDATE petitions
        #     SET content = %s::jsonb,
        #         updated_at = CURRENT_TIMESTAMP
        #     WHERE tenant_id = %s AND id = %s
        #     """,
        #     (json.dumps(content2), str(tenant_id), int(petition_id)),
        #     fetch=None,
        # )

    # -----------------------------
    # DOCX export
    # -----------------------------
    def export_docx_bytes(self, tenant_id: str, petition_id: int) -> bytes:
        """
        Gera um DOCX em memória.
        Prioridade do conteúdo:
        1) última versão gerada (petition_versions.rendered_content)
        2) content.texto_livre
        3) fallback: render_from_template não chamado ainda -> gera texto mínimo
        """
        from docx import Document  # python-docx

        petition = self.get_petition(tenant_id, petition_id)
        if not petition:
            raise ValueError("Petição não encontrada")

        versions = self.list_versions(tenant_id, petition_id)
        latest = versions[0] if versions else None

        if latest and latest.get("rendered_content"):
            text = _safe_str(latest.get("rendered_content"))
        else:
            content = self._merge_content(petition.get("content"))
            tl = content.get("texto_livre")

            # se texto_livre vier “vazado” como dict, tenta recuperar
            if isinstance(tl, dict):
                tl = tl.get("texto_livre", "")

            text = (_safe_str(tl)).strip()
            if not text:
                # fallback: um “mini-render” para não exportar vazio
                ptype = self._norm_petition_type(petition.get("petition_type") or "")
                ctx = self.preparar_contexto_peticao(str(petition.get("process_id")), tenant_id) or {}
                text = self._jinja.from_string(self._builtin_template_for(ptype)).render(**ctx, content=content).strip()

        doc = Document()

        # Título simples
        title = f"Petição - {str(petition.get('petition_type') or '').upper()}  (ID {petition_id})"
        doc.add_heading(title, level=1)

        # Corpo preservando quebras
        for line in text.splitlines():
            if line.strip() == "":
                doc.add_paragraph("")  # mantém “linha em branco”
            else:
                doc.add_paragraph(line)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
